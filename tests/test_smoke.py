"""Smoke tests — verify the package imports, the config loads, and the CLI works.

These tests do NOT call the Anthropic API or NotebookLM. They check:
- every module imports cleanly
- pydantic config validates the sample yaml
- regex parsers in verify/internal handle representative inputs
- the JSON-decision-file validator catches the malformed-JSON gotcha
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest


def test_top_level_imports():
    import thesis_generator
    assert hasattr(thesis_generator, "__version__")
    assert hasattr(thesis_generator, "ThesisProject")
    assert hasattr(thesis_generator, "run_pipeline")


def test_all_submodules_import():
    """Every module must import without side effects (no auto-API-calls etc.)."""
    import thesis_generator.cli  # noqa: F401
    import thesis_generator.config  # noqa: F401
    import thesis_generator.env_check  # noqa: F401
    import thesis_generator.inventory  # noqa: F401
    import thesis_generator.pipeline  # noqa: F401
    import thesis_generator.writer  # noqa: F401
    import thesis_generator.docx_ops  # noqa: F401
    import thesis_generator.reports  # noqa: F401
    import thesis_generator.orchestration  # noqa: F401
    import thesis_generator.notebooklm  # noqa: F401
    import thesis_generator.review  # noqa: F401
    import thesis_generator.verify  # noqa: F401


def test_sample_yaml_validates():
    """The sample project's thesis.yaml must validate against the pydantic model."""
    from thesis_generator.config import ThesisProject

    sample = Path(__file__).resolve().parent.parent / "examples" / "sample_project"
    project = ThesisProject.load(sample)
    assert project.title.startswith("Word-of-Mouth")
    assert project.language.value == "pl"
    assert project.citation_style.value == "apa7"
    assert project.notebooklm is not None
    assert "REPLACE-WITH-YOUR-UUID" in project.notebooklm.library_url
    assert "A" in project.pipeline.verification_rings


def test_humanize_policy_defaults():
    from thesis_generator.config import HumanizationPolicy

    h = HumanizationPolicy()
    assert h.remove_em_dashes is True
    assert "triangulacja" in h.forbidden_words


def test_ring_a_citation_regex():
    """Brief gotcha: the parens-citation regex must handle Polish surnames + et al."""
    from thesis_generator.verify.internal import _CITATION_PARENS

    text = "Konsumenci ufają (Hocha i Janicka, 2002, s. 137) bardziej niż reklamie (Trusov et al., 2009)."
    matches = list(_CITATION_PARENS.finditer(text))
    assert len(matches) == 2
    assert matches[0].group(2) == "2002"
    assert matches[1].group(2) == "2009"


def test_json_decision_file_validator_catches_malformed():
    """The string-escape gotcha (#8): a malformed JSON file must fail before we apply edits."""
    from thesis_generator.orchestration.patterns import assert_well_formed_decision_file

    bad = Path("_tmp_bad.json")
    bad.write_text('{"edits": [{"para": "not an int", "old": "x"}]}', encoding="utf-8")
    try:
        with pytest.raises((TypeError, ValueError)):
            assert_well_formed_decision_file(str(bad))
    finally:
        bad.unlink(missing_ok=True)


def test_json_decision_file_validator_passes_valid():
    from thesis_generator.orchestration.patterns import assert_well_formed_decision_file

    good = Path("_tmp_good.json")
    good.write_text(
        json.dumps({"edits": [{"para": 5, "old": "x", "new": "y"}]}),
        encoding="utf-8",
    )
    try:
        assert_well_formed_decision_file(str(good))  # must not raise
    finally:
        good.unlink(missing_ok=True)


def test_progress_report_renders():
    from thesis_generator.reports.progress import ProgressReport

    report = ProgressReport(
        phase="Writer — sekcja 4.1",
        status="IN_PROGRESS",
        icon="✍️",
        file_metrics={"Paragrafy": 132, "Znaków": 45_000},
        completed=["wczytano draft", "wygenerowano inventory"],
        next_steps=["uruchom ring A", "uruchom ring B"],
        blockers=[],
    )
    md = report.render()
    assert "Writer — sekcja 4.1" in md
    assert "IN_PROGRESS" in md
    assert "Paragrafy" in md
    assert "brak" in md  # default blockers note


def test_cli_help_works():
    """Typer must import and the help text must render."""
    from typer.testing import CliRunner

    from thesis_generator.cli import app

    runner = CliRunner()
    result = runner.invoke(app, ["--help"])
    assert result.exit_code == 0
    assert "thesis-generator" in result.stdout
    assert "verify-env" in result.stdout
    assert "notebooklm" in result.stdout


def test_humanize_em_dash_replacement():
    """The user said 3x: replace ` — ` with `, `. Test the core transform."""
    from thesis_generator.config import HumanizationPolicy
    from thesis_generator.docx_ops.humanize import HumanizationStats, _humanize_text

    policy = HumanizationPolicy()
    stats = HumanizationStats()
    out = _humanize_text("Marketing — szczególnie WoM — ma znaczenie.", policy, stats)
    assert " — " not in out
    assert stats.em_dashes_replaced == 2


def test_inventory_author_year_filename_parser():
    from thesis_generator.inventory import _guess_author_year_from_filename

    cases = [
        ("Hoch_2002_Product_Experience.pdf", ("Hoch", "2002")),
        ("Trusov_et_al_2009_WoM_Effects.pdf", ("Trusov", "2009")),
        ("Babbie_2013_Podstawy_badań_społecznych.pdf", ("Babbie", "2013")),
    ]
    for fname, expected in cases:
        author, year = _guess_author_year_from_filename(fname)
        assert (author, year) == expected, f"{fname} → ({author}, {year}) != {expected}"


# ----- Regression tests for bugs caught by the e2e run on 2026-05-23 -----


def test_ring_a_handles_polish_et_al():
    """Regression: original regex missed `(Trusov i in., 2009)` and `Goldstein i in. (2008)`.

    E2E test on a stub draft revealed both parens and narrative forms were
    skipping Polish "i in." (= English "et al."), leading to one false-positive
    orphan and one missed cited-but-missing-from-bib citation.
    """
    from thesis_generator.verify.internal import _CITATION_NARRATIVE, _CITATION_PARENS

    parens_text = "eWOM kształtuje sprzedaż (Trusov i in., 2009, s. 90)."
    assert _CITATION_PARENS.search(parens_text), "Parens form with 'i in.' must match"

    narrative_text = "Goldstein i in. (2008) pokazali coś."
    m = _CITATION_NARRATIVE.search(narrative_text)
    assert m is not None, "Narrative form with 'i in.' must match"
    assert m.group(2) == "2008"


def test_notebooklm_strict_verdict_parser_no_false_positives():
    """Regression: original parser matched stray 'OK' inside prose, giving false-positive verdicts.

    E2E run reported Trusov_2009 as `OK` even though NotebookLM returned
    free-form academic prose without any VERDICT line. The strict parser
    requires `^VERDICT:` at the start of a line.
    """
    from thesis_generator.notebooklm.adapter import NotebookLMAdapter, VerificationVerdict

    # Stray "OK" in prose must NOT trigger a verdict
    prose = "Tak, OK to ważna kwestia. Marketing szeptany jest istotny (Hoch 2002)."
    assert NotebookLMAdapter._parse_verdict(prose) == VerificationVerdict.UNKNOWN

    # Genuine VERDICT line at start MUST trigger
    genuine = "Po analizie źródła:\nVERDICT: OK\nEXCERPT: \"Konsumenci ufają WoM.\""
    assert NotebookLMAdapter._parse_verdict(genuine) == VerificationVerdict.OK

    # VERDICT mid-line must NOT trigger (avoids matching "Odpowiedź to VERDICT: ...")
    midline = "Odpowiedź to VERDICT: OK ale w prozie"
    assert NotebookLMAdapter._parse_verdict(midline) == VerificationVerdict.UNKNOWN


def test_repair_docx_handles_ctrl_f9_disaster(tmp_path):
    """Synthetic Ctrl+F9 corruption — repair must restore body text + remove fake shell."""
    import zipfile
    import shutil

    from docx import Document

    from thesis_generator.docx_ops.repair import diagnose, repair_docx

    # Build a corrupted docx by wrapping a normal body in a fake fldChar shell.
    base = tmp_path / "base.docx"
    Document().save(str(base))  # python-docx writes a minimal valid skeleton
    with zipfile.ZipFile(base, "r") as zin:
        items = {n: zin.read(n) for n in zin.namelist()}

    # Replace document.xml with a corrupted body: 1 fake fldChar pair wrapping
    # 150 instrText elements, no w:t at all.
    body_inner = "".join(
        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:r><w:instrText> akapit numer {i} treść badawcza</w:instrText></w:r>'
        f"</w:p>"
        for i in range(150)
    )
    corrupted_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        '<w:p><w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> FAKEFIELD </w:instrText></w:r></w:p>'
        f"{body_inner}"
        '<w:p><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
        "</w:body></w:document>"
    ).encode("utf-8")
    items["word/document.xml"] = corrupted_xml
    corrupted = tmp_path / "corrupted.docx"
    with zipfile.ZipFile(corrupted, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in items.items():
            zout.writestr(n, data)

    # Diagnose must flag it
    diag = diagnose(corrupted)
    assert any("Ctrl+F9" in d for d in diag.diagnosed), f"Diagnose missed corruption: {diag.diagnosed}"

    # Repair must fix it
    report = repair_docx(corrupted)
    assert report.ok, f"Repair did not declare OK: {report.needs_manual}"
    assert any("Pass A" in f for f in report.fixes_applied)
    assert any("Pass B" in f for f in report.fixes_applied)
    assert report.backup_path and report.backup_path.exists()

    # Post-repair: real text is back
    doc = Document(str(corrupted))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "akapit numer 0" in body_text
    assert "akapit numer 149" in body_text


def test_effective_path_helpers_autodiscover_conventional_layout(tmp_path):
    """ThesisProject.effective_*() helpers must find files at conventional paths."""
    from docx import Document

    from thesis_generator.config import ThesisProject

    # Scaffold the conventional layout without using tg init (so we test the
    # discovery logic directly).
    proj = tmp_path / "proj"
    (proj / "inputs" / "sources").mkdir(parents=True)
    (proj / "inputs" / "research_data" / "surveys").mkdir(parents=True)
    (proj / "inputs" / "research_data" / "interviews").mkdir(parents=True)
    (proj / "inputs" / "school").mkdir(parents=True)
    # Drop conventional files
    Document().save(str(proj / "inputs" / "draft.docx"))
    Document().save(str(proj / "inputs" / "school" / "regulation.docx"))
    (proj / "inputs" / "research_data" / "surveys" / "responses.xlsx").write_bytes(b"PK\x03\x04")
    (proj / "inputs" / "research_data" / "interviews" / "idi01.pdf").write_bytes(b"%PDF-1.4")

    (proj / "thesis.yaml").write_text(
        """title: T
author: A
inputs:
  draft: inputs/draft.docx
  sources_dir: inputs/sources
notebooklm:
  library_url: https://notebooklm.google.com/notebook/abc
  library_name: x
""",
        encoding="utf-8",
    )
    p = ThesisProject.load(proj)

    assert p.effective_research_data_dir() == proj / "inputs" / "research_data"
    assert p.effective_interviews_dir() == proj / "inputs" / "research_data" / "interviews"
    assert p.effective_school_dir() == proj / "inputs" / "school"
    assert p.effective_regulation() == proj / "inputs" / "school" / "regulation.docx"
    files = p.effective_research_data_files()
    assert any(f.name == "responses.xlsx" for f in files)


def test_visual_marker_parser_handles_three_kinds():
    """Parser must handle TABELA, WYKRES, ILUSTRACJA + SUGEROWANY plus polish/english keys."""
    from thesis_generator.visuals.markers import VisualKind, parse_markers

    text = """
    Tu jest jakiś akapit z teorią (Kahneman, 2011, s. 45).
    [TABELA 1: Charakterystyka próby][Źródło: opracowanie własne][Dane: research_data/surveys/responses.xlsx sheet=Metryczka cells=A1:C20]
    Następnie pojawia się obraz.
    [ILUSTRACJA 2: Sensorama (1962)][Źródło: Wikimedia Commons][Plik: visuals/sensorama.jpg][Szerokość: 11cm]
    A potem wykres.
    [WYKRES 3: Rozkład ocen Q22][Source: opracowanie własne][Dane: data.xlsx sheet=H1 cells=A1:C5][Type: line]
    Na koniec sugestia:
    [SUGEROWANY WYKRES: Wzrost rynku VR 2018-2023][Opis: do uzupełnienia]
    """
    markers = parse_markers(text)
    assert len(markers) == 4

    kinds = [m.kind for m in markers]
    assert kinds == [VisualKind.TABLE, VisualKind.IMAGE, VisualKind.CHART, VisualKind.SUGGESTED]

    # TABELA — data spec parsed
    t = markers[0]
    assert t.title == "Charakterystyka próby"
    assert t.source == "opracowanie własne"
    assert t.data is not None
    assert t.data.sheet == "Metryczka"
    assert t.data.cells == "A1:C20"
    assert "responses.xlsx" in t.data.file

    # ILUSTRACJA — width parsed
    i = markers[1]
    assert i.file == "visuals/sensorama.jpg"
    assert i.width_cm == 11.0

    # WYKRES — chart_type via English "Type:"
    w = markers[2]
    assert w.chart_type == "line"

    # SUGEROWANY — description captured, no data
    s = markers[3]
    assert s.is_suggestion
    assert "do uzupełnienia" in s.description


def test_visuals_registry_assigns_sequential_numbers():
    """Numbering follows document order, separate sequences for tables vs figures."""
    from thesis_generator.visuals.markers import VisualKind, VisualMarker
    from thesis_generator.visuals.registry import VisualsRegistry

    reg = VisualsRegistry()
    markers = [
        VisualMarker(kind=VisualKind.TABLE, declared_number=None, title="T-A"),
        VisualMarker(kind=VisualKind.CHART, declared_number=None, title="C-A"),
        VisualMarker(kind=VisualKind.TABLE, declared_number=None, title="T-B"),
        VisualMarker(kind=VisualKind.IMAGE, declared_number=None, title="I-A"),
        VisualMarker(kind=VisualKind.CHART, declared_number=None, title="C-B"),
    ]
    entries = reg.assign(markers)
    # Tables: T-A=1, T-B=2; Figures: C-A=1, I-A=2, C-B=3
    nums = [(e.kind.value, e.number) for e in entries]
    assert nums == [
        ("TABELA", 1), ("WYKRES", 1), ("TABELA", 2), ("ILUSTRACJA", 2), ("WYKRES", 3),
    ]
    assert "Tabela 1. T-A" in reg.render_spis_tabel()
    assert "Rysunek 3. C-B" in reg.render_spis_rysunkow()


def test_chart_rendering_writes_png(tmp_path):
    """Smoke: render a tiny bar chart to PNG — proves matplotlib path works."""
    from thesis_generator.visuals.charts import ChartData, render_chart

    data = ChartData(
        title="Test bar",
        x_labels=["A", "B", "C"],
        series=[("series1", [1.0, 2.5, 3.7])],
        y_label="Y",
        chart_type="bar",
    )
    out = tmp_path / "test.png"
    rendered = render_chart(data, out)
    assert rendered == out
    assert out.exists()
    assert out.stat().st_size > 1000  # non-trivial PNG


def test_table_styling_uniform(tmp_path):
    """Apply uniform table styling — borders, header shading, font normalized."""
    from docx import Document

    from thesis_generator.visuals.tables import style_all_tables

    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "Header A"
    tbl.cell(0, 1).text = "Header B"
    tbl.cell(1, 0).text = "row1"
    tbl.cell(1, 1).text = "row2"
    n = style_all_tables(doc)
    assert n == 1

    p = tmp_path / "styled.docx"
    doc.save(str(p))
    # Re-open and inspect: header cell must have D9D9D9 shading
    from docx.oxml.ns import qn
    d2 = Document(str(p))
    hdr = d2.tables[0].cell(0, 0)
    tcPr = hdr._tc.find(qn("w:tcPr"))
    shd = tcPr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")).upper() == "D9D9D9"


def test_notebook_writer_prompt_includes_required_constraints():
    """Per-section prompt must include the proven-session constraints (no em-dash, APA, intro authors, no meta)."""
    from thesis_generator.notebooklm.writer import NotebookSectionSpec, build_writer_prompt

    spec = NotebookSectionSpec(
        id="1.2.3",
        title="Marketing szeptany online",
        focus="Skupić się na eWOM, Trusov 2009, Rosario 2016.",
        priority_sources=["Trusov_2009", "Rosario_2016"],
        target_chars=6000,
    )
    prompt = build_writer_prompt(spec)
    assert "1.2.3" in prompt
    assert "Marketing szeptany online" in prompt
    assert "APA 7" in prompt
    assert "em-dashy" in prompt.lower() or "em-dash" in prompt.lower()
    assert "Trusov_2009" in prompt
    assert "Bibliografia" in prompt


def test_notebook_response_parser_splits_body_from_bibliography():
    """parse_notebook_response must separate prose from `Bibliografia wykorzystana` section."""
    from thesis_generator.notebooklm.writer import parse_notebook_response

    raw = (
        "Marketing szeptany ma długą historię w literaturze (Hoch, 2002, s. 137). "
        "Trusov i in. (2009, s. 90) wykazali, że eWOM przewyższa reklamę.\n\n"
        "Bibliografia wykorzystana w podrozdziale\n\n"
        "Hoch, S. J. (2002). Product experience is seductive. Journal of Consumer Research, 29(3), 448-454.\n\n"
        "Trusov, M., Bucklin, R. E., & Pauwels, K. (2009). Effects of WoM. Journal of Marketing, 73(5), 90-102."
    )
    body, bib, trunc = parse_notebook_response(raw)
    assert "Hoch" in body and "Trusov" in body
    assert "Bibliografia wykorzystana" not in body
    assert len(bib) == 2
    assert "Hoch, S. J." in bib[0]
    assert trunc is False


def test_correction_prompt_lists_fixes_in_structured_format():
    """build_correction_prompt must include every fix with FLAG + correction hint."""
    from thesis_generator.notebooklm.correction import CitationFix, build_correction_prompt

    fixes = [
        CitationFix(
            section_id="1.2",
            paragraph_tag="P0042",
            original_sentence="Konsumenci ufają WoM bardziej niż reklamie.",
            cited_as="(Hoch, 2002, s. 137)",
            flag="⚠️ STRONA",
            correction_hint="actual claim is on s. 142, not 137",
        ),
        CitationFix(
            section_id="1.2",
            paragraph_tag="P0058",
            original_sentence="Rosario 2016 dowodzi, że produkty doświadczalne są wyjątkiem.",
            cited_as="(Rosario, 2016, s. 305)",
            flag="❌ TREŚĆ",
            correction_hint="WYMAGA DECYZJI — Rosario nie pisze o produktach doświadczalnych",
        ),
    ]
    prompt = build_correction_prompt("1.2", fixes)
    assert "1.2" in prompt
    assert "⚠️ STRONA" in prompt
    assert "❌ TREŚĆ" in prompt
    assert "s. 137" in prompt
    assert "WYMAGA DECYZJI" in prompt
    assert "PRZED:" in prompt
    assert "PO:" in prompt


def test_corrections_parser_extracts_before_after():
    """parse_corrections must extract structured (before, after, comment) from NotebookLM reply."""
    from thesis_generator.notebooklm.correction import parse_corrections

    raw = '''#### Poprawka 1. [P0042]
PRZED: "Konsumenci ufają WoM bardziej niż reklamie."
PO: "Hoch (2002, s. 142) wykazał, że konsumenci uznają rekomendacje za bardziej wiarygodne niż reklamy marek."
KOMENTARZ: Strona była błędna — claim znajduje się na s. 142, nie 137.

#### Poprawka 2. [P0058]
PRZED: "Rosario 2016 dowodzi, że produkty doświadczalne są wyjątkiem."
PO: "Rosario i in. (2016, s. 305) wskazali, że eWOM ma silniejszy efekt w produktach doświadczalnych."
KOMENTARZ: Poprawiono błędną parafrazę.
'''
    out = parse_corrections(raw)
    assert len(out) == 2
    assert out[0]["paragraph_tag"] == "P0042"
    assert "s. 142" in out[0]["after"]
    assert out[1]["paragraph_tag"] == "P0058"


def test_citation_audit_adapter_handles_missing_skill(monkeypatch, tmp_path):
    """Adapter must raise a clear error when the skill isn't installed."""
    from thesis_generator.citation_audit.adapter import CitationAuditError, _default_skill_dir
    import shutil

    if (Path.home() / ".claude" / "skills" / "thesis-citation-audit").exists():
        # If real skill exists on this machine, just ensure _default_skill_dir returns its path
        p = _default_skill_dir()
        assert p.exists()
    else:
        # Otherwise, missing skill must raise with helpful message
        import pytest

        with pytest.raises(CitationAuditError) as exc_info:
            _default_skill_dir()
        assert "thesis-citation-audit" in str(exc_info.value)


def test_notebook_pipeline_importable_and_runs_steps_in_order():
    """Pipeline module must import cleanly and expose the public entry point."""
    from thesis_generator import pipeline_notebook

    assert hasattr(pipeline_notebook, "run_notebook_pipeline")
    assert hasattr(pipeline_notebook, "NotebookPipelineResult")


def test_numeric_claim_extractor_finds_stats_and_percentages(tmp_path):
    """Extractor must catch M=, %, χ², t, p, d, n, and exclude bare year numbers."""
    from docx import Document

    from thesis_generator.numbers_audit.extractor import extract_numeric_claims

    doc = Document()
    doc.add_paragraph("Wstęp z roku 2020 nie powinien być traktowany jako twierdzenie liczbowe.")
    doc.add_paragraph(
        "Analiza pokazała M = 4.21 (SD = 1.13) dla grupy posiadaczy (n = 62). "
        "Test χ² = 12.3 dał p < 0.001 z efektem d = 0.42."
    )
    doc.add_paragraph(
        "65% respondentów oceniło rekomendację pozytywnie, a 6 respondentów wskazało odpowiedź neutralną."
    )
    p = tmp_path / "test.docx"
    doc.save(str(p))

    claims = extract_numeric_claims(p)
    # First paragraph (pure year) should NOT produce a claim
    assert not any(c.paragraph_idx == 0 for c in claims)
    # Second paragraph should have multiple stats
    second = [c for c in claims if c.paragraph_idx == 1]
    assert second, "Stats sentence not extracted"
    found = " ".join(second[0].numbers_found)
    assert "M = 4.21" in found.replace(" =", "=").replace("= ", "=") or "M=4.21" in found.replace(" ", "")
    # Third paragraph: percentage + respondent count
    third = [c for c in claims if c.paragraph_idx == 2]
    assert third
    assert any("%" in n for n in third[0].numbers_found)


def test_hypothesis_extractor_finds_H_declarations_and_verdicts(tmp_path):
    """Extractor must catch H1/H2/H3 + their verdict keywords."""
    from docx import Document

    from thesis_generator.numbers_audit.extractor import extract_hypotheses

    doc = Document()
    doc.add_paragraph(
        "H1: Konsumenci ufają rekomendacjom znajomych bardziej niż reklamie. "
        "H2: Demo VR zwiększa intencję zakupową."
    )
    doc.add_paragraph(
        "Wyniki testu t Welcha wskazują, że H1 została potwierdzona (p < 0.001). "
        "Hipoteza H2 została częściowo potwierdzona, jedynie w grupie posiadaczy."
    )
    p = tmp_path / "h.docx"
    doc.save(str(p))

    hyps = extract_hypotheses(p)
    ids = {h.hypothesis_id for h in hyps}
    assert "H1" in ids
    assert "H2" in ids
    # Verdict keywords are detected
    verdicts = [h.verdict_keyword for h in hyps if h.verdict_keyword]
    assert any("potwierdzon" in v for v in verdicts)


def test_recompute_diff_parser_handles_jsonl():
    """recomputer._parse_opus_diffs must parse JSON-lines from the DIFFS block."""
    from thesis_generator.numbers_audit.recomputer import _parse_opus_diffs

    response = '''## SUMMARY
4 OK, 1 MISMATCH found. The headset count is off by 2.

## DIFFS
{"paragraph_tag": "P0012", "claimed": "65%", "recomputed": "64.7%", "delta": "Δ=0.3pp", "status": "OK", "note": "Rounding."}
{"paragraph_tag": "P0042", "claimed": "6 respondentów", "recomputed": "4 respondentów", "delta": "Δ=2", "status": "MISMATCH", "note": "Author counted wrong rows."}
{"paragraph_tag": "P0058", "claimed": "M = 4.21", "recomputed": "M = 4.21", "delta": "exact", "status": "OK", "note": "Confirmed."}
'''
    diffs = _parse_opus_diffs(response)
    assert len(diffs) == 3
    mismatches = [d for d in diffs if d.status == "MISMATCH"]
    assert len(mismatches) == 1
    assert mismatches[0].paragraph_tag == "P0042"
    assert "4 respondentów" in mismatches[0].recomputed


def test_hypothesis_verdict_parser_handles_opus_output():
    """consistency._parse_verdicts must extract VERDICT blocks."""
    from thesis_generator.numbers_audit.consistency import _parse_verdicts

    response = '''#### VERDICT
HYPOTHESIS_ID: H1
DECLARATION: "Konsumenci ufają WoM bardziej niż reklamie."
CLAIMED_VERDICT: potwierdzona
JUDGED_VERDICT: SUPPORTED
REASONING: Test t Welcha z p < 0.001 i efektem d = 0.62 mocno wspiera hipotezę. Recomputed M_wom = 4.16 vs M_reklama = 1.63.
RELEVANT_STATS: P0042, P0058

#### VERDICT
HYPOTHESIS_ID: H2
DECLARATION: "Demo VR zwiększa intencję zakupową."
CLAIMED_VERDICT: potwierdzona
JUDGED_VERDICT: OVER_INTERPRETED
REASONING: Recomputed efekt d = 0.18 jest mały. Autor traktuje to jako mocne potwierdzenie, ale to tylko trend.
RELEVANT_STATS: P0072
'''
    verdicts = _parse_verdicts(response)
    assert len(verdicts) == 2
    assert verdicts[0].hypothesis_id == "H1"
    assert verdicts[0].judged_verdict == "SUPPORTED"
    assert verdicts[1].judged_verdict == "OVER_INTERPRETED"
    assert "P0042" in verdicts[0].relevant_stats


def test_audit_only_runs_shallow_on_existing_thesis(tmp_path):
    """Audit pipeline must run shallow (no LLM) on a real docx and never mutate it."""
    from docx import Document

    from thesis_generator.audit_only import run_audit
    from thesis_generator.config import ThesisProject

    # Scaffold a minimal project with a draft that has planted Ring A bugs
    proj = tmp_path / "audit-test"
    (proj / "inputs" / "sources").mkdir(parents=True)
    (proj / "inputs" / "research_data" / "surveys").mkdir(parents=True)

    doc = Document()
    doc.add_heading("Wstęp", level=1)
    doc.add_paragraph(
        "Marketing szeptany ma długą historię (Hoch, 2002, s. 137). "
        "Trusov i in. (2009, s. 90) wykazali — że eWOM przewyższa reklamę."
    )
    doc.add_heading("BIBLIOGRAFIA", level=1)
    doc.add_paragraph("Hoch, S. J. (2002). Product experience. Journal of Consumer Research, 29(3), 448-454.")
    doc.add_paragraph("Trusov, M. et al. (2009). Effects of WoM. Journal of Marketing, 73(5), 90-102.")
    doc.add_paragraph("Krawczyk, A. (2018). Polskie pokolenie Z.")  # orphan
    draft_path = proj / "inputs" / "draft.docx"
    doc.save(str(draft_path))
    original_size = draft_path.stat().st_size
    original_mtime = draft_path.stat().st_mtime

    (proj / "thesis.yaml").write_text(
        """title: Test
author: T
inputs:
  draft: inputs/draft.docx
  sources_dir: inputs/sources
  raw_data: []
notebooklm:
  library_url: https://notebooklm.google.com/notebook/x
  library_name: Test
""",
        encoding="utf-8",
    )
    project = ThesisProject.load(proj)

    # Shallow audit — no LLM calls, no swarm, no reviewer
    result = run_audit(project, deep=False)

    # Read-only invariant: draft must be untouched
    assert draft_path.stat().st_size == original_size, "Audit mutated draft size"
    assert draft_path.stat().st_mtime == original_mtime, "Audit mutated draft mtime"

    # Phases that ran
    assert result.env is not None
    assert result.ring_a is not None
    assert result.visuals is not None
    assert result.humanize is not None
    assert result.data_audit is not None

    # Deep-only phases skipped
    assert result.citation_audit is None
    assert result.reviewer is None
    assert result.fixes is None

    # Planted bugs detected
    assert any(s == "Krawczyk" for s, y in [(e.first_surname, e.year) for e in result.ring_a.orphaned_in_bib])
    assert result.ring_a.em_dash_count >= 1  # planted ` — `

    # Consolidated report written
    assert result.consolidated_report is not None
    assert result.consolidated_report.exists()
    body = result.consolidated_report.read_text(encoding="utf-8")
    assert "THESIS AUDIT" in body
    assert "Ring A" in body
    assert "Krawczyk" in body  # orphan surfaced in report


def test_setup_wizard_imports_and_non_interactive_run(tmp_path, monkeypatch):
    """The setup wizard must import cleanly and scaffold a project in non-interactive mode."""
    from thesis_generator.wizard import run_setup_wizard

    monkeypatch.setenv("ANTHROPIC_API_KEY", "")  # ensure we go through CLI path
    project = tmp_path / "wiz-test"
    code = run_setup_wizard(project, non_interactive=True)
    # Even if some env_check warnings fire (no claude on CI), wizard returns 0
    # for soft failures. The thesis.yaml must exist regardless.
    assert (project / "thesis.yaml").exists(), "thesis.yaml not written"
    assert (project / "inputs" / "research_data" / "surveys" / "README.md").exists()
    assert (project / "inputs" / "visuals" / "README.md").exists()
    assert (project / "inputs" / "school" / "README.md").exists()
    # Re-run must be idempotent — no error, files still there
    code2 = run_setup_wizard(project, non_interactive=True)
    assert (project / "thesis.yaml").exists()


def test_notebooklm_grounding_prompt_forces_named_source():
    """The improved grounding prompt must force NotebookLM to use ONE source.

    E2E showed earlier prompts let NotebookLM ramble about adjacent topics.
    The fix enforces source-specific quoting.
    """
    from thesis_generator.notebooklm.adapter import NotebookLMAdapter

    q = NotebookLMAdapter._build_grounding_question(
        source_ref="Hoch 2002",
        paraphrase="Konsumenci ufają WoM.",
        cited_page="s. 137",
    )
    assert "WYŁĄCZNIE" in q  # must constrain to one source
    assert "Hoch 2002" in q
    assert "NIE używaj żadnych innych źródeł" in q
    assert "s. 137" in q


def test_claude_cli_adapter_imports_and_handles_missing_binary():
    """ClaudeCLI must import cleanly and produce a clear error when `claude` is absent.

    We don't make a live API call here (CI shouldn't depend on subscription auth).
    Just verify the type is constructible when claude is on PATH and raises a
    useful error when it isn't.
    """
    from thesis_generator.llm import ClaudeCLI, ClaudeCLIError
    import shutil

    if shutil.which("claude"):
        cli = ClaudeCLI()
        assert cli.claude_bin
        assert cli.default_cwd.exists()
    else:
        try:
            ClaudeCLI()
        except ClaudeCLIError as e:
            assert "claude" in str(e).lower()
        else:
            raise AssertionError("Expected ClaudeCLIError when `claude` is missing")


def test_claude_cli_build_cmd_includes_expected_flags():
    """The generated `claude -p` invocation must include the safety/isolation flags."""
    from thesis_generator.llm.claude_cli import ClaudeCLI
    import shutil

    if not shutil.which("claude"):
        import pytest

        pytest.skip("`claude` not on PATH")
    cli = ClaudeCLI()
    cmd, _env = cli._build_cmd(
        model="haiku",
        system="be brief",
        allowed_tools=None,
        json_schema=None,
        max_budget_usd=0.5,
    )
    assert "--print" in cmd
    assert "--no-session-persistence" in cmd
    assert "--permission-mode" in cmd
    assert "bypassPermissions" in cmd
    assert "--model" in cmd
    assert "haiku" in cmd
    assert "--system-prompt" in cmd
    assert "--max-budget-usd" in cmd
    assert "0.5" in cmd
    assert "--tools" in cmd  # no allowed_tools → tools disabled


def test_humanize_forbidden_words_handles_declensions():
    """Regression: 'triangulacja' check missed 'triangulacji' (Polish genitive).

    The original `if w.lower() in body_text.lower()` required exact substring
    match, which fails for any inflected form. Fixed by stem-matching the
    first ~6 chars and using a `\\b<stem>\\w*` regex.
    """
    from thesis_generator.config import HumanizationPolicy
    from thesis_generator.docx_ops.humanize import HumanizationStats
    from docx import Document
    import tempfile, os

    policy = HumanizationPolicy()  # default forbidden_words includes "triangulacja"
    doc = Document()
    doc.add_paragraph("Badanie zastosowało metodologię opartą na triangulacji danych.")
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(path)
        from thesis_generator.docx_ops.humanize import humanize_docx
        stats = humanize_docx(__import__("pathlib").Path(path), policy)
        joined = " ".join(stats.forbidden_words_found)
        assert "triangulacja" in joined, f"Expected stem match on 'triangulacji'; got: {stats.forbidden_words_found}"
    finally:
        os.unlink(path)
