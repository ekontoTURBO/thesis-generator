"""Generate a 3-page stub thesis draft for end-to-end pipeline testing.

Content: short version of the user's real thesis topic — VR commerce + WoM.
Citations: pulled from the brief's "key sources" list (Hoch, Trusov, Rosario,
Berger, Batorski) — all confirmed to be in the NotebookLM library.

Deliberate test plants:
- 1 em-dash (` — `) → Ring A + humanize should flag/fix
- 1 orphan in bibliography (Krawczyk) → Ring A should catch
- 1 missing-from-bibliography citation (Goldstein) → Ring A should catch
- 1 forbidden word ("triangulacja") → humanize should flag
- Mix of (Author, Year, s. X) citations matching APA 7 → Ring C will verify
"""

from docx import Document
from docx.shared import Pt

doc = Document()

# Set default font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Title
title = doc.add_heading("Word-of-Mouth marketing w wirtualnej rzeczywistości — wstęp", level=1)

# Section 1
doc.add_heading("1. Marketing szeptany jako siła napędowa decyzji konsumenckich", level=2)
doc.add_paragraph(
    "Marketing szeptany od dekad pozostaje jednym z najsilniejszych mechanizmów "
    "wpływu społecznego w decyzjach zakupowych. Badania empiryczne pokazują, że konsumenci "
    "uznają rekomendacje innych konsumentów za bardziej wiarygodne niż komunikaty marketingowe "
    "marek (Hoch, 2002, s. 137). Wraz z rozwojem mediów cyfrowych mechanika ta przeniosła się "
    "do środowiska online, gdzie eWOM kształtuje sprzedaż w sposób mierzalny i trwały "
    "(Trusov i in., 2009, s. 90). Meta-analiza Rosario (2016, s. 305) potwierdza, że efekt "
    "eWOM jest silniejszy w kategoriach produktów doświadczalnych niż użytkowych — to obserwacja "
    "kluczowa dla VR jako medium ze swojej natury doświadczalnego."
)

# Section 2
doc.add_heading("2. Wirtualna rzeczywistość jako nowe medium doświadczenia produktowego", level=2)
doc.add_paragraph(
    "Wirtualna rzeczywistość redefiniuje pojęcie doświadczenia produktowego, pozwalając "
    "konsumentowi wejść w interakcję z produktem zanim podejmie decyzję zakupową. Mechanika "
    "wiralności w VR opiera się na trzech elementach: zaskoczeniu, emocjach i społecznej "
    "wymianie (Berger, 2013, s. 14). Badania polskich konsumentów online wskazują, że "
    "intencja podzielenia się doświadczeniem rośnie proporcjonalnie do intensywności "
    "zaangażowania zmysłowego (Batorski, 2006, s. 103). Dodatkowo Goldstein i in. (2008) "
    "pokazali, że normatywne sygnały społeczne w środowisku cyfrowym znacząco zmieniają "
    "intencje zakupowe."
)

# Section 3 — methodology snippet
doc.add_heading("3. Metodologia badania", level=2)
doc.add_paragraph(
    "Badanie zastosowało metodologię mieszaną opartą na triangulacji danych ilościowych "
    "(CAWI, N=123) i jakościowych (IDI, N=3). Próba została dobrana metodą warstwową "
    "z uwzględnieniem podziału na posiadaczy i nie-posiadaczy gogli VR. Analiza statystyczna "
    "obejmowała testy χ², t Welcha oraz korelacje Pearsona, z poziomem istotności p < 0,05."
)

# Bibliography section
doc.add_heading("BIBLIOGRAFIA", level=1)
doc.add_heading("I. Literatura naukowa", level=2)
bib = [
    "Batorski, D., & Olcoń-Kubicka, M. (2006). Prowadzenie badań przez Internet — podstawowe zagadnienia metodologiczne. Studia Socjologiczne, 3(182), 99-132.",
    "Berger, J. (2013). Contagious: Why Things Catch On. Simon & Schuster.",
    "Hoch, S. J. (2002). Product experience is seductive. Journal of Consumer Research, 29(3), 448-454.",
    "Krawczyk, A. (2018). Polskie pokolenie Z a media społecznościowe. Wydawnictwo Naukowe PWN.",  # orphan plant
    "Rosario, A. B., Sotgiu, F., De Valck, K., & Bijmolt, T. H. A. (2016). The effect of electronic word of mouth on sales: A meta-analytic review of platform, product, and metric factors. Journal of Marketing Research, 53(3), 297-318.",
    "Trusov, M., Bucklin, R. E., & Pauwels, K. (2009). Effects of word-of-mouth versus traditional marketing: Findings from an internet social networking site. Journal of Marketing, 73(5), 90-102.",
]
for entry in bib:
    doc.add_paragraph(entry)

out = "C:/Users/erykc/Desktop/thesis-generator/_e2e_test/inputs/draft.docx"
doc.save(out)
print(f"Wrote {out} ({len(doc.paragraphs)} paragraphs)")
