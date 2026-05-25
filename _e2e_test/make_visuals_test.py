"""Live e2e for the visuals pipeline.

Builds:
  - inputs/research_data/surveys/sample.xlsx with one sheet of survey-like data
  - inputs/visuals/sample_image.png (a generated placeholder)
  - inputs/draft_visuals_test.docx with 3 markers (TABELA + WYKRES + ILUSTRACJA)

Then runs `process_visuals` and prints diagnostics.
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from openpyxl import Workbook

PROJECT = Path(__file__).parent

# --- 1. Build a sample xlsx -------------------------------------------------

surveys_dir = PROJECT / "inputs" / "research_data" / "surveys"
surveys_dir.mkdir(parents=True, exist_ok=True)
xlsx_path = surveys_dir / "sample.xlsx"

wb = Workbook()
ws = wb.active
ws.title = "Q22_Q23"
ws["A1"] = "Ocena"
ws["B1"] = "Q22 (zaufanie WoM)"
ws["C1"] = "Q23 (zaufanie reklama)"
ratings = ["1", "2", "3", "4", "5"]
q22 = [1, 1, 19, 55, 29]
q23 = [24, 49, 22, 4, 6]
for i, r in enumerate(ratings, start=2):
    ws.cell(row=i, column=1, value=r)
    ws.cell(row=i, column=2, value=q22[i - 2])
    ws.cell(row=i, column=3, value=q23[i - 2])

# A small metryczka table
ws2 = wb.create_sheet("Metryczka")
ws2["A1"] = "Cecha"
ws2["B1"] = "n"
ws2["C1"] = "%"
ws2["A2"] = "Kobiety"
ws2["B2"] = 62
ws2["C2"] = 51.7
ws2["A3"] = "Mężczyźni"
ws2["B3"] = 58
ws2["C3"] = 48.3
ws2["A4"] = "Razem"
ws2["B4"] = 120
ws2["C4"] = 100.0

wb.save(xlsx_path)
print(f"✓ Wrote {xlsx_path}")

# --- 2. Build a placeholder image -------------------------------------------

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

visuals_dir = PROJECT / "inputs" / "visuals"
visuals_dir.mkdir(parents=True, exist_ok=True)
img_path = visuals_dir / "sample_image.png"
fig, ax = plt.subplots(figsize=(4, 3))
ax.text(0.5, 0.5, "Sample diagram\n(VR commerce funnel)", ha="center", va="center", fontsize=14)
ax.axis("off")
plt.savefig(str(img_path), dpi=120, bbox_inches="tight")
plt.close()
print(f"✓ Wrote {img_path}")

# --- 3. Build a stub draft with 3 markers -----------------------------------

doc = Document()
doc.add_heading("Test sekcji z wizualizacjami", level=1)

doc.add_paragraph(
    "Badanie ankietowe (CAWI, N=120) obejmowało dwie grupy respondentów. "
    "Charakterystyka próby przedstawia rozkład płci w badaniu."
)
doc.add_paragraph(
    "[TABELA 1: Charakterystyka próby badawczej][Źródło: opracowanie własne na podstawie wyników ankiety CAWI]"
    "[Dane: inputs/research_data/surveys/sample.xlsx sheet=Metryczka cells=A1:C4]"
)

doc.add_paragraph(
    "Analiza rozkładów ocen w odpowiedziach Q22 (zaufanie do WoM) i Q23 (zaufanie do reklam) "
    "pokazuje wyraźną asymetrię między obiema źródłami informacji marketingowych."
)
doc.add_paragraph(
    "[WYKRES 1: Rozkład ocen Q22 vs Q23 (skala 1-5)][Źródło: opracowanie własne]"
    "[Dane: inputs/research_data/surveys/sample.xlsx sheet=Q22_Q23 cells=A1:C6][Typ: bar]"
)

doc.add_paragraph(
    "Schemat lejka marketingowego w środowisku VR obrazuje kolejne etapy podejmowania decyzji "
    "przez konsumenta."
)
doc.add_paragraph(
    "[ILUSTRACJA 1: Lejek marketingowy w VR commerce][Źródło: opracowanie własne]"
    "[Plik: inputs/visuals/sample_image.png][Szerokość: 11cm]"
)

doc.add_paragraph(
    "Zaprezentowane dane wskazują na istotną przewagę WoM nad reklamą tradycyjną w kategorii produktów "
    "doświadczalnych, jakimi są gogle VR."
)
doc.add_heading("BIBLIOGRAFIA", level=1)
doc.add_paragraph("(stub bibliography)")

draft_path = PROJECT / "inputs" / "draft_visuals_test.docx"
doc.save(str(draft_path))
print(f"✓ Wrote {draft_path}")

# --- 4. Run the visuals pipeline against it ---------------------------------

# Patch thesis.yaml to point at the new draft for this run
import yaml
yaml_path = PROJECT / "thesis.yaml"
cfg = yaml.safe_load(yaml_path.read_text(encoding="utf-8"))
original_draft = cfg["inputs"]["draft"]
cfg["inputs"]["draft"] = "inputs/draft_visuals_test.docx"
yaml_path.write_text(yaml.dump(cfg, allow_unicode=True, sort_keys=False), encoding="utf-8")

try:
    from thesis_generator.config import ThesisProject
    from thesis_generator.visuals import process_visuals

    project = ThesisProject.load(PROJECT)
    output = PROJECT / "output" / "draft_with_visuals.docx"
    output.parent.mkdir(exist_ok=True)
    result = process_visuals(project, output_path=output)

    print()
    print("=" * 60)
    print(f"INSERTED: {result.inserted}")
    print(f"SKIPPED:  {len(result.skipped)}")
    for title, reason in result.skipped:
        print(f"   - {title}: {reason}")
    if result.registry:
        print(f"Tables in registry:   {len(result.registry.tables)}")
        print(f"Figures in registry:  {len(result.registry.figures)}")
    for n in result.notes:
        print(f"NOTE: {n}")
    print()
    print(f"Output: {output}")

    # Verify counts in the saved file
    d = Document(str(output))
    print(f"Final: {len(d.paragraphs)} paragraphs, {len(d.tables)} tables, {len(d.inline_shapes)} images")
finally:
    # Restore original draft path
    cfg["inputs"]["draft"] = original_draft
    yaml_path.write_text(yaml.dump(cfg, allow_unicode=True, sort_keys=False), encoding="utf-8")
