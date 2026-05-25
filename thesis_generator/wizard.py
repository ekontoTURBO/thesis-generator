"""Interactive first-run setup wizard.

The point of `tg setup` is to remove every friction step a first-time user
would otherwise hit: no Anthropic API key (use `claude` CLI), no manual
NotebookLM auth dance (offer to run it), no hand-editing of thesis.yaml,
no guesswork about what folder is for what.

Idempotent — safe to re-run. If a thesis.yaml already exists, the wizard
respects it and only updates fields the user explicitly changes.
"""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

import typer
import yaml
from rich.console import Console
from rich.panel import Panel
from rich.prompt import Confirm, Prompt
from rich.table import Table


@dataclass(slots=True)
class WizardState:
    project_dir: Path
    is_existing_project: bool
    existing_config: dict | None = None


# Stable section header that walks the user through the steps.
_STEPS = 7


def _step_header(console: Console, n: int, title: str) -> None:
    console.print()
    console.rule(f"[bold cyan][{n}/{_STEPS}] {title}[/bold cyan]")


def run_setup_wizard(
    project_dir: Path | None = None,
    *,
    console: Console | None = None,
    non_interactive: bool = False,
) -> int:
    """Run the first-run wizard. Returns process-style exit code (0 = success)."""
    console = console or Console()

    console.print()
    console.print(
        Panel.fit(
            "[bold]thesis-generator — first-run setup[/bold]\n\n"
            "Walks you through every step. Safe to re-run.",
            border_style="cyan",
        )
    )

    # -------- Step 1: system checks --------
    _step_header(console, 1, "System checks")
    ok = _system_checks(console)
    if not ok:
        console.print(
            "\n[red]One or more system requirements are missing.[/red] "
            "Fix the items above and re-run [bold]tg setup[/bold]."
        )
        return 1

    # -------- Step 2: pick project directory --------
    _step_header(console, 2, "Project location")
    state = _resolve_project_dir(console, project_dir, non_interactive=non_interactive)
    if state is None:
        return 1

    # -------- Step 3: thesis basics (title/author/etc.) --------
    _step_header(console, 3, "Thesis basics")
    basics = _collect_basics(console, state, non_interactive=non_interactive)

    # -------- Step 4: NotebookLM library --------
    _step_header(console, 4, "NotebookLM library")
    notebooklm = _resolve_notebooklm(console, state, non_interactive=non_interactive)

    # -------- Step 5: scaffold folder structure --------
    _step_header(console, 5, "Folder structure")
    _scaffold(console, state, basics, notebooklm)

    # -------- Step 6: starter draft / explainer --------
    _step_header(console, 6, "Starter files")
    _explain_inputs(console, state)

    # -------- Step 7: final env_check --------
    _step_header(console, 7, "Final check")
    final_ok = _run_final_check(console, state.project_dir)

    # -------- Wrap-up --------
    console.print()
    if final_ok:
        console.print(
            Panel.fit(
                f"[bold green]✓ Setup complete[/bold green]\n\n"
                f"Project: [cyan]{state.project_dir}[/cyan]\n"
                f"Config:  [cyan]{state.project_dir / 'thesis.yaml'}[/cyan]",
                border_style="green",
            )
        )
        _print_next_steps(console, state.project_dir)
        return 0
    else:
        console.print(
            Panel.fit(
                "[bold yellow]Setup completed with warnings.[/bold yellow]\n"
                "Read the env_check output above and fix what's flagged, then re-run [bold]tg verify-env[/bold].",
                border_style="yellow",
            )
        )
        return 0  # not a hard failure


# ============================================================================
# Step 1 — system checks
# ============================================================================


def _system_checks(console: Console) -> bool:
    tbl = Table(show_header=False, box=None, padding=(0, 1))
    tbl.add_column(); tbl.add_column(); tbl.add_column()
    all_ok = True

    # Python version
    py_ok = sys.version_info >= (3, 11)
    tbl.add_row(
        "✓" if py_ok else "✗",
        "Python",
        f"{sys.version.split()[0]}" + ("" if py_ok else " [red](need 3.11+)[/red]"),
    )
    all_ok = all_ok and py_ok

    # claude CLI
    claude_bin = shutil.which("claude")
    if claude_bin:
        tbl.add_row("✓", "claude CLI", claude_bin)
        # Auth check
        try:
            from thesis_generator.llm import ClaudeCLI
            ok, msg = ClaudeCLI(claude_bin=claude_bin).check_auth()
            tbl.add_row("✓" if ok else "✗", "  → auth", msg if ok else f"[red]{msg}[/red]")
            all_ok = all_ok and ok
        except Exception as e:
            tbl.add_row("✗", "  → auth", f"[red]check failed: {e}[/red]")
            all_ok = False
    else:
        api = os.environ.get("ANTHROPIC_API_KEY")
        if api:
            tbl.add_row("✓", "Anthropic SDK", f"ANTHROPIC_API_KEY set ({len(api)} chars)")
        else:
            tbl.add_row(
                "✗",
                "claude CLI",
                "[red]not on PATH[/red] — install from https://claude.com/code, then run `claude /login`",
            )
            all_ok = False

    # NotebookLM skill
    skill_dir = Path.home() / ".claude" / "skills" / "notebooklm"
    if skill_dir.exists():
        tbl.add_row("✓", "NotebookLM skill", str(skill_dir))
        # Auth
        try:
            proc = subprocess.run(
                [sys.executable, "scripts/run.py", "auth_manager.py", "status"],
                cwd=str(skill_dir),
                env={**os.environ, "PYTHONIOENCODING": "utf-8"},
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=60,
                check=False,
            )
            out = (proc.stdout or "").lower()
            authed = "authenticated" in out or "✓" in (proc.stdout or "")
            stale = "may need re-authentication" in (proc.stdout or "")
            if authed:
                tbl.add_row("✓", "  → auth", "authenticated" + (" [yellow](stale, may expire soon)[/yellow]" if stale else ""))
            else:
                tbl.add_row("○", "  → auth", "[yellow]not authenticated[/yellow] — wizard will offer to fix")
        except Exception as e:
            tbl.add_row("○", "  → auth", f"[yellow]check skipped: {type(e).__name__}[/yellow]")
    else:
        tbl.add_row(
            "○",
            "NotebookLM skill",
            "[yellow]not installed[/yellow] — install from https://github.com/PleasePrompto/notebooklm-skill",
        )
        # Not strictly required if user opts out of Ring C, so don't fail.

    console.print(tbl)
    return all_ok


# ============================================================================
# Step 2 — project location
# ============================================================================


def _resolve_project_dir(
    console: Console, project_dir: Path | None, *, non_interactive: bool
) -> WizardState | None:
    if project_dir is None:
        if non_interactive:
            console.print("[red]Project path required in non-interactive mode.[/red]")
            return None
        default = str(Path.cwd() / "my-thesis")
        answer = Prompt.ask("Where should this thesis project live?", default=default)
        project_dir = Path(answer).expanduser().resolve()
    else:
        project_dir = Path(project_dir).expanduser().resolve()

    existing = (project_dir / "thesis.yaml").exists()
    if existing:
        console.print(f"[yellow]Existing project found at[/yellow] {project_dir}")
        if not non_interactive:
            cont = Confirm.ask("Re-run setup on it (update fields, keep your files)?", default=True)
            if not cont:
                console.print("Aborted.")
                return None
        try:
            cfg = yaml.safe_load((project_dir / "thesis.yaml").read_text(encoding="utf-8"))
        except Exception:
            cfg = None
        return WizardState(project_dir=project_dir, is_existing_project=True, existing_config=cfg)

    console.print(f"[green]New project will be created at[/green] {project_dir}")
    return WizardState(project_dir=project_dir, is_existing_project=False)


# ============================================================================
# Step 3 — thesis basics
# ============================================================================


def _collect_basics(console: Console, state: WizardState, *, non_interactive: bool) -> dict:
    cfg = state.existing_config or {}
    defaults = {
        "title": cfg.get("title", "Tytuł pracy licencjackiej"),
        "author": cfg.get("author", "Imię Nazwisko"),
        "school": cfg.get("school", ""),
        "promotor": cfg.get("promotor", ""),
        "language": cfg.get("language", "pl"),
        "citation_style": cfg.get("citation_style", "apa7"),
    }
    if non_interactive:
        return defaults

    console.print("[dim]Press Enter to keep the default in brackets.[/dim]\n")
    return {
        "title": Prompt.ask("Thesis title", default=defaults["title"]),
        "author": Prompt.ask("Author", default=defaults["author"]),
        "school": Prompt.ask("School (optional)", default=defaults["school"]),
        "promotor": Prompt.ask("Promotor (optional)", default=defaults["promotor"]),
        "language": Prompt.ask("Language", choices=["pl", "en"], default=defaults["language"]),
        "citation_style": Prompt.ask(
            "Citation style", choices=["apa7", "vancouver", "chicago"], default=defaults["citation_style"]
        ),
    }


# ============================================================================
# Step 4 — NotebookLM library
# ============================================================================


def _resolve_notebooklm(console: Console, state: WizardState, *, non_interactive: bool) -> dict | None:
    cfg = state.existing_config or {}
    existing = cfg.get("notebooklm") if isinstance(cfg.get("notebooklm"), dict) else None

    # Try to read libraries the NotebookLM skill already knows about.
    library_json = Path.home() / ".claude" / "skills" / "notebooklm" / "data" / "library.json"
    known: list[tuple[str, str]] = []  # (name, url)
    if library_json.exists():
        try:
            data = json.loads(library_json.read_text(encoding="utf-8"))
            for n in (data.get("notebooks") or {}).values():
                if n.get("url"):
                    known.append((n.get("name", "(no name)"), n["url"]))
        except Exception:
            pass

    if known:
        console.print(f"[green]Found {len(known)} NotebookLM library(ies)[/green] already configured:")
        for i, (name, url) in enumerate(known, 1):
            console.print(f"  [bold]{i}.[/bold] {name}")
            console.print(f"     [dim]{url}[/dim]")

    if non_interactive:
        if existing:
            return existing
        if known:
            return {"library_url": known[0][1], "library_name": known[0][0]}
        return None

    if existing:
        console.print(
            f"[yellow]Existing config points at:[/yellow] {existing.get('library_name', '(unnamed)')} "
            f"— [dim]{existing.get('library_url', '')}[/dim]"
        )
        keep = Confirm.ask("Keep this library?", default=True)
        if keep:
            return existing

    # Pick from known or paste URL or skip
    options = [str(i) for i in range(1, len(known) + 1)] + ["paste", "skip"]
    default = "1" if known else "paste"
    choice = Prompt.ask(
        f"Pick a library [{'/'.join(options)}]",
        choices=options,
        default=default,
    )
    if choice == "skip":
        console.print("[yellow]Skipping NotebookLM[/yellow] — Ring C verification will be disabled.")
        return None
    if choice == "paste":
        url = Prompt.ask("Paste NotebookLM URL (https://notebooklm.google.com/notebook/<UUID>)")
        name = Prompt.ask("Short name for this library", default="My thesis library")
        return {"library_url": url.strip(), "library_name": name.strip()}

    idx = int(choice) - 1
    selected = {"library_url": known[idx][1], "library_name": known[idx][0]}

    # System-prompt-reminder — critical for thesis-quality output.
    console.print()
    console.print(
        Panel(
            "[bold yellow]⚠️ One-time required step:[/bold yellow] set the [bold]System Instructions[/bold] in this notebook.\n\n"
            "1. Open the notebook in NotebookLM\n"
            "2. Click [bold]Settings[/bold] (gear icon) → [bold]System Instructions[/bold]\n"
            "3. Paste the contents of [cyan]docs/notebooklm_system_prompt.pl.pdf[/cyan]\n"
            "4. Save\n\n"
            "Without this, NotebookLM gives conversational answers instead of APA-formatted "
            "thesis text. The per-section prompts that `tg notebook-write` sends are short "
            "BECAUSE the global behavior lives in this system prompt.",
            border_style="yellow",
            title="NotebookLM System Prompt",
        )
    )
    return selected


# ============================================================================
# Step 5 — scaffold the project (delegates to init logic)
# ============================================================================


def _scaffold(
    console: Console,
    state: WizardState,
    basics: dict,
    notebooklm: dict | None,
) -> None:
    project_dir = state.project_dir
    project_dir.mkdir(parents=True, exist_ok=True)

    # Top-level tool dirs
    for sub in ("_reports", "_state", "output"):
        (project_dir / sub).mkdir(exist_ok=True)

    # inputs/ tree (reuse CLI helpers so README content stays in one place)
    from thesis_generator.cli import (
        _readme_inputs,
        _readme_notes,
        _readme_research_data,
        _readme_research_existing,
        _readme_research_interviews,
        _readme_research_observations,
        _readme_research_surveys,
        _readme_school,
        _readme_sources,
        _readme_visuals,
    )

    tree = {
        "inputs": _readme_inputs(),
        "inputs/sources": _readme_sources(),
        "inputs/research_data": _readme_research_data(),
        "inputs/research_data/surveys": _readme_research_surveys(),
        "inputs/research_data/interviews": _readme_research_interviews(),
        "inputs/research_data/observations": _readme_research_observations(),
        "inputs/research_data/existing": _readme_research_existing(),
        "inputs/visuals": _readme_visuals(),
        "inputs/school": _readme_school(),
        "inputs/notes": _readme_notes(),
    }
    created = 0
    for rel, readme in tree.items():
        folder = project_dir / rel
        if not folder.exists():
            folder.mkdir(parents=True)
            created += 1
        rp = folder / "README.md"
        if not rp.exists():
            rp.write_text(readme, encoding="utf-8")
    console.print(f"  ✓ {len(tree)} input folders ({created} newly created)")
    console.print(f"  ✓ README.md placed in each folder explaining what goes there")

    # Placeholder draft if missing
    draft_path = project_dir / "inputs" / "draft.docx"
    if not draft_path.exists():
        try:
            from docx import Document
            d = Document()
            d.add_heading("(Replace this stub with your actual thesis draft)", level=1)
            d.add_paragraph(
                "This is a placeholder created by `tg setup`. Replace it with your real .docx, "
                "or drop your existing draft in this location with the same filename."
            )
            d.save(str(draft_path))
            console.print("  ✓ placeholder inputs/draft.docx")
        except Exception:
            draft_path.touch()

    # Write thesis.yaml — merging with existing config if present
    yaml_path = project_dir / "thesis.yaml"
    new_cfg = state.existing_config or {}
    new_cfg.update(basics)
    new_cfg.setdefault("inputs", {})
    new_cfg["inputs"]["draft"] = "inputs/draft.docx"
    new_cfg["inputs"]["sources_dir"] = "inputs/sources"
    new_cfg["inputs"].setdefault("research_data_dir", "inputs/research_data")
    new_cfg["inputs"].setdefault("school_dir", "inputs/school")
    new_cfg["inputs"].setdefault("raw_data", [])

    if notebooklm:
        new_cfg["notebooklm"] = {
            "library_url": notebooklm["library_url"],
            "library_name": notebooklm["library_name"],
            "max_words_per_query": 400,
            "parallel_queries": 3,
            "timeout_seconds": 300,
        }

    new_cfg.setdefault("models", {
        "writer": "claude-opus-4-7",
        "auditor": "claude-opus-4-7",
        "reviewer": "claude-opus-4-7",
        "rewriter": "claude-sonnet-4-6",
        "verifier_per_source": "claude-haiku-4-5-20251001",
    })
    new_cfg.setdefault("humanization", {
        "remove_em_dashes": True,
        "forbidden_words": ["triangulacja", "rygorystycznie", "holistycznie", "synergiczny", "paradygmat"],
    })
    new_cfg.setdefault("output", {"dir": "output", "docx_name": "Praca.docx", "pdf": False, "versioned": True})
    new_cfg.setdefault("pipeline", {
        "verification_rings": ["A", "B", "C"] if notebooklm else ["A", "B"],
        "max_parallel_haiku": 12,
        "fail_on_grade_below": 4.0,
    })

    yaml_path.write_text(
        yaml.dump(new_cfg, allow_unicode=True, sort_keys=False, default_flow_style=False),
        encoding="utf-8",
    )
    console.print(f"  ✓ thesis.yaml written ({len(new_cfg)} top-level sections)")


# ============================================================================
# Step 6 — explainer
# ============================================================================


def _explain_inputs(console: Console, state: WizardState) -> None:
    p = state.project_dir
    tbl = Table(show_header=True, header_style="bold cyan")
    tbl.add_column("Drop here", style="cyan")
    tbl.add_column("What goes there")
    tbl.add_row(f"{p / 'inputs' / 'draft.docx'}", "Your existing partial thesis (replace the placeholder)")
    tbl.add_row(f"{p / 'inputs' / 'sources'}/", "Academic PDFs you cite (Author_Year_Title.pdf)")
    tbl.add_row(f"{p / 'inputs' / 'research_data' / 'surveys'}/", "Survey responses (.xlsx, .csv)")
    tbl.add_row(f"{p / 'inputs' / 'research_data' / 'interviews'}/", "IDI/FGI transcripts (.pdf, .docx)")
    tbl.add_row(f"{p / 'inputs' / 'research_data' / 'existing'}/", "Secondary data (GUS, reports as xlsx)")
    tbl.add_row(f"{p / 'inputs' / 'visuals'}/", "Ready images for ILUSTRACJA markers")
    tbl.add_row(f"{p / 'inputs' / 'school'}/", "regulation.docx (formatting rules), brief.pdf")
    console.print(tbl)


# ============================================================================
# Step 7 — final env_check
# ============================================================================


def _run_final_check(console: Console, project_dir: Path) -> bool:
    console.print("Running [bold]tg verify-env[/bold]…\n")
    try:
        from thesis_generator.config import ThesisProject
        from thesis_generator.env_check import check_env

        project = ThesisProject.load(project_dir)
        result = check_env(project)
        console.print(result.render())
        return result.ok
    except Exception as e:
        console.print(f"[red]env_check raised: {e}[/red]")
        return False


def _print_next_steps(console: Console, project_dir: Path) -> None:
    console.print("\n[bold]Next steps:[/bold]")
    console.print(f"  1. Replace [cyan]{project_dir / 'inputs' / 'draft.docx'}[/cyan] with your real draft")
    console.print(f"  2. Drop your source PDFs into [cyan]{project_dir / 'inputs' / 'sources'}/[/cyan]")
    console.print(f"  3. Drop research data into [cyan]{project_dir / 'inputs' / 'research_data'}/[/cyan] (surveys/, interviews/, etc.)")
    console.print(f"  4. (Re-)run [green]tg verify-env {project_dir}[/green] when you've added the real files")
    console.print(f"  5. [green]tg verify {project_dir} --rings A[/green]  — fast internal audit")
    console.print(f"  6. [green]tg run {project_dir}[/green]               — full pipeline\n")
