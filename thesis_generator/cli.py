"""CLI entry point — `thesis-generator` / `tg` commands.

Subcommands:
    init         — scaffold a new project folder with a starter thesis.yaml
    verify-env   — run environment check only
    inventory    — index sources + extract style fingerprint
    verify       — run verification rings (A, B, C selectable)
    review       — run independent reviewer pass
    audit-data   — recompute statistics from raw Excel
    humanize     — strip em-dashes, flag forbidden words
    run          — run the whole pipeline end-to-end
    notebooklm   — manage NotebookLM auth / libraries
"""

from __future__ import annotations

import asyncio
from pathlib import Path

import typer
from rich.console import Console
from rich.table import Table

from thesis_generator import __version__
from thesis_generator.config import ThesisProject

app = typer.Typer(
    name="thesis-generator",
    help="Verify-first multi-agent thesis pipeline. Built on the proven workflow from real bachelor's-thesis sessions.",
    no_args_is_help=True,
    add_completion=False,
)
console = Console()


@app.command("setup")
def setup_cmd(
    project_dir: Path = typer.Argument(None, help="Where to create / re-configure the project (will prompt if omitted)"),
    non_interactive: bool = typer.Option(False, "--non-interactive", "-y", help="Skip prompts, accept defaults"),
) -> None:
    """Guided first-run setup. Walks you through everything — system check, NotebookLM auth, project scaffold, thesis.yaml fields, final env_check."""
    from thesis_generator.wizard import run_setup_wizard

    code = run_setup_wizard(project_dir, console=console, non_interactive=non_interactive)
    if code != 0:
        raise typer.Exit(code)


@app.command("init")
def init_cmd(
    project_dir: Path = typer.Argument(..., help="Where to scaffold the project"),
    title: str = typer.Option("Tytuł pracy licencjackiej", help="Thesis title"),
    author: str = typer.Option("Imię Nazwisko", help="Author"),
) -> None:
    """Scaffold a new project folder with the full inputs/ tree + thesis.yaml."""
    project_dir = project_dir.expanduser().resolve()
    project_dir.mkdir(parents=True, exist_ok=True)

    yaml_path = project_dir / "thesis.yaml"
    if yaml_path.exists():
        console.print(f"[yellow]thesis.yaml already exists at {yaml_path}[/yellow]")
        raise typer.Exit(1)

    yaml_path.write_text(_starter_yaml(title, author), encoding="utf-8")

    # Top-level tool dirs
    for sub in ("_reports", "_state", "output"):
        (project_dir / sub).mkdir(exist_ok=True)

    # Full inputs/ tree with README in each folder explaining what goes there.
    tree = {
        "inputs": _readme_inputs(),
        "inputs/sources": _readme_sources(),
        "inputs/research_data": _readme_research_data(),
        "inputs/research_data/surveys": _readme_research_surveys(),
        "inputs/research_data/interviews": _readme_research_interviews(),
        "inputs/research_data/observations": _readme_research_observations(),
        "inputs/research_data/existing": _readme_research_existing(),
        "inputs/visuals": _readme_visuals(),
        "inputs/school": _readme_school(),
        "inputs/notes": _readme_notes(),
    }
    for rel, readme in tree.items():
        folder = project_dir / rel
        folder.mkdir(parents=True, exist_ok=True)
        readme_path = folder / "README.md"
        if not readme_path.exists():
            readme_path.write_text(readme, encoding="utf-8")

    # Drop a placeholder draft so verify-env doesn't immediately fail.
    placeholder = project_dir / "inputs" / "draft.docx"
    if not placeholder.exists():
        try:
            from docx import Document
            d = Document()
            d.add_heading("(Replace this stub with your actual thesis draft)", level=1)
            d.add_paragraph("This is a placeholder created by `tg init`. Replace it with your real .docx.")
            d.save(str(placeholder))
        except Exception:
            placeholder.touch()  # last-resort empty file

    console.print(f"[green]✓[/green] Initialized project at {project_dir}")
    console.print()
    console.print("[bold]Next steps:[/bold]")
    console.print(f"  1. Replace [cyan]{project_dir}/inputs/draft.docx[/cyan] with your real draft")
    console.print(f"  2. Drop your source PDFs into [cyan]{project_dir}/inputs/sources/[/cyan]")
    console.print(f"  3. Drop research data into the appropriate [cyan]inputs/research_data/[/cyan] subfolder")
    console.print(f"  4. Edit [cyan]{yaml_path}[/cyan] — set the NotebookLM library URL")
    console.print(f"  5. Run [green]tg verify-env {project_dir}[/green]")


@app.command("verify-env")
def verify_env_cmd(project_dir: Path) -> None:
    """Run the environment check (hard gate before any other step)."""
    from thesis_generator.env_check import check_env

    project = ThesisProject.load(project_dir)
    result = check_env(project)
    console.print(result.render())
    raise typer.Exit(0 if result.ok else 1)


@app.command("inventory")
def inventory_cmd(project_dir: Path) -> None:
    """Index source PDFs and extract style fingerprint from the existing draft."""
    from thesis_generator.inventory import build_inventory

    project = ThesisProject.load(project_dir)
    inv = build_inventory(project)
    table = Table(title=f"Inventory — {len(inv.sources)} sources")
    table.add_column("Author"); table.add_column("Year"); table.add_column("Size"); table.add_column("Path")
    for s in inv.sources[:30]:
        table.add_row(s.author_guess[:20], s.year_guess or "—", f"{s.size_bytes // 1024} KB", Path(s.path).name)
    if len(inv.sources) > 30:
        table.add_row("…", "…", "…", f"+ {len(inv.sources) - 30} more")
    console.print(table)
    for n in inv.notes:
        console.print(f"[dim]• {n}[/dim]")


@app.command("verify")
def verify_cmd(
    project_dir: Path,
    rings: str = typer.Option("A,B,C", help="Comma-separated rings to run: A, B, C"),
) -> None:
    """Run verification rings."""
    from thesis_generator.verify.internal import run_ring_a, save_report as save_a
    from thesis_generator.verify.haiku_per_source import run_ring_b, save_report as save_b
    from thesis_generator.verify.notebooklm_ring import run_ring_c, save_report as save_c

    project = ThesisProject.load(project_dir)
    selected = {r.strip().upper() for r in rings.split(",")}

    if "A" in selected:
        result = run_ring_a(project)
        path = save_a(result, project)
        console.print(f"[green]Ring A[/green] → {path} | {'PASS' if result.passed else 'FAIL'}")

    if "B" in selected:
        result = asyncio.run(run_ring_b(project))
        path = save_b(result, project)
        console.print(f"[green]Ring B[/green] → {path} | {'PASS' if result.passed else 'FAIL'}")

    if "C" in selected:
        if project.notebooklm is None:
            console.print("[yellow]Ring C skipped: no notebooklm in thesis.yaml[/yellow]")
        else:
            console.print("[dim]Ring C needs a paraphrase map — invoke via `tg run` or programmatic API for now.[/dim]")


@app.command("review")
def review_cmd(project_dir: Path) -> None:
    """Run the independent reviewer (zero-context Opus, strict professor)."""
    from thesis_generator.review.independent import run_independent_review, save_report

    project = ThesisProject.load(project_dir)
    console.print("[dim]Running independent reviewer (one Opus call, no tool use, may take 30-60s)…[/dim]")
    result = run_independent_review(project)
    path = save_report(result, project)
    console.print(f"[bold]Grade:[/bold] {result.grade}")
    console.print(f"Report: {path}")


@app.command("audit-data")
def audit_data_cmd(project_dir: Path) -> None:
    """Recompute every statistic in the draft from raw Excel."""
    from thesis_generator.verify.data_audit import run_data_audit, save_report

    project = ThesisProject.load(project_dir)
    result = run_data_audit(project)
    path = save_report(result, project)
    console.print(f"Data audit → {path}")
    console.print(f"  matches: {len(result.matches)} | likely mismatches: {len(result.likely_off_by)}")


@app.command("humanize")
def humanize_cmd(project_dir: Path) -> None:
    """Strip em-dashes and flag forbidden buzzwords."""
    from thesis_generator.docx_ops.humanize import humanize_docx

    project = ThesisProject.load(project_dir)
    draft = project.resolve_input(project.inputs.draft)
    stats = humanize_docx(draft, project.humanization)
    console.print(stats.render())


@app.command("recompute-data")
def recompute_data_cmd(
    project_dir: Path,
    max_claims: int = typer.Option(80, help="Cap on number of claims to verify (cost control)"),
) -> None:
    """Independent Opus recomputation of every numeric claim. Produces a parallel xlsx workbook."""
    from thesis_generator.numbers_audit import run_recompute

    project = ThesisProject.load(project_dir)
    console.print("[dim]Dispatching Opus auditor (may take 2-5 min — writes its own Python, runs it via Bash, computes from raw data)…[/dim]")
    result = run_recompute(project, max_claims=max_claims)
    console.print(f"[bold]Total claims checked:[/bold] {result.total}")
    console.print(f"  [green]OK[/green]: {result.ok_count}")
    console.print(f"  [red]Mismatches[/red]: {result.mismatches}")
    console.print(f"  [yellow]Unverifiable[/yellow]: {result.unverifiable}")
    if result.workbook_path:
        console.print(f"Parallel workbook: [cyan]{result.workbook_path}[/cyan]")
    if result.markdown_report_path:
        console.print(f"Report: [cyan]{result.markdown_report_path}[/cyan]")
    for n in result.notes:
        console.print(f"[dim]{n}[/dim]")


@app.command("check-hypotheses")
def check_hypotheses_cmd(project_dir: Path) -> None:
    """Run recompute + judge each H1/H2/H3 against the independently recomputed numbers."""
    from thesis_generator.numbers_audit import check_hypothesis_consistency, run_recompute

    project = ThesisProject.load(project_dir)
    console.print("[dim]Step 1: Opus recomputes statistics…[/dim]")
    rec = run_recompute(project)
    console.print(f"[dim]→ {rec.total} claims checked ({rec.mismatches} mismatches)[/dim]\n")
    console.print("[dim]Step 2: Opus judges per-hypothesis consistency…[/dim]")
    cons = check_hypothesis_consistency(project, rec)
    console.print(f"[bold]Hypotheses judged:[/bold] {len(cons.verdicts)}")
    for v in cons.verdicts:
        icon = {"SUPPORTED": "✅", "NOT_SUPPORTED": "❌", "PARTIALLY": "⚠️",
                "OVER_INTERPRETED": "🚨", "INSUFFICIENT_DATA": "❓"}.get(v.judged_verdict, "•")
        console.print(f"  {icon} {v.hypothesis_id}: {v.judged_verdict}")
    if cons.markdown_report_path:
        console.print(f"Report: [cyan]{cons.markdown_report_path}[/cyan]")


@app.command("audit")
def audit_cmd(
    project_dir: Path,
    deep: bool = typer.Option(True, "--deep/--shallow", help="Deep = include Haiku swarm + Opus reviewer (slow). Shallow = only regex/data checks."),
    skip: str = typer.Option("", help="Comma-separated phases to skip: env,inventory,ring_a,citation_audit,orchestrate,data,reviewer,visuals,humanize"),
) -> None:
    """Audit an EXISTING thesis without writing or mutating it. Read-only verification + consolidated report."""
    from thesis_generator.audit_only import run_audit

    project = ThesisProject.load(project_dir)
    skipped = tuple(s.strip() for s in skip.split(",") if s.strip())
    console.print(f"[dim]Running audit (deep={deep}), skipping: {skipped or 'nothing'}…[/dim]\n")
    result = run_audit(project, deep=deep, skip=skipped)

    console.print()
    console.rule("[bold cyan]AUDIT SUMMARY[/bold cyan]")
    if result.reviewer and result.reviewer.grade is not None:
        console.print(f"[bold]Reviewer grade:[/bold] {result.reviewer.grade}")
    if result.ring_a:
        status = "[green]PASS[/green]" if result.ring_a.passed else "[red]FAIL[/red]"
        console.print(
            f"[bold]Ring A:[/bold] {status} — "
            f"{len(result.ring_a.missing_in_bib)} missing-from-bib, "
            f"{len(result.ring_a.orphaned_in_bib)} orphans, "
            f"{result.ring_a.em_dash_count} em-dashes"
        )
    if result.citation_audit:
        console.print(
            f"[bold]Citation audit:[/bold] {result.citation_audit.total_problems} problems / "
            f"{result.citation_audit.total_citations} checked"
        )
    if result.fixes:
        console.print(f"[bold]Suggested fixes (Opus):[/bold] {len(result.fixes.fixes)}")
    if result.data_audit:
        console.print(
            f"[bold]Data audit:[/bold] {len(result.data_audit.likely_off_by)} likely mismatches"
        )
    if result.visuals:
        v = result.visuals
        console.print(
            f"[bold]Visuals:[/bold] {v.total_existing_tables} tables, "
            f"{v.tables + v.charts + v.images} markers"
        )
    if result.humanize:
        console.print(
            f"[bold]AI tells:[/bold] {result.humanize.em_dashes_replaced} em-dashes, "
            f"{len(result.humanize.forbidden_words_found)} buzzwords"
        )
    if result.skipped:
        console.print(f"[yellow]Phases skipped:[/yellow] {[p for p, _ in result.skipped]}")
    console.print()
    console.print(f"[bold green]Consolidated report:[/bold green] {result.consolidated_report}")


@app.command("notebook-write")
def notebook_write_cmd(
    project_dir: Path,
    section_id: str = typer.Option(..., "--section", "-s", help="e.g. 1.2 or 4.1"),
    title: str = typer.Option(..., "--title", "-t", help="Subsection title"),
    focus: str = typer.Option(..., "--focus", "-f", help="2-5 sentences describing scope"),
    priority_sources: str = typer.Option("", "--sources", help="Comma-separated author-year refs"),
    target_chars: int = typer.Option(6000, help="Target length"),
) -> None:
    """STEP 1 — Ask NotebookLM to write one subsection step-by-step (uses the system prompt configured in NotebookLM)."""
    from thesis_generator.notebooklm import NotebookLMAdapter, NotebookSectionSpec, write_section

    project = ThesisProject.load(project_dir)
    if project.notebooklm is None:
        console.print("[red]No notebooklm block in thesis.yaml[/red]")
        raise typer.Exit(1)

    adapter = NotebookLMAdapter(
        library_url=project.notebooklm.library_url,
        library_name=project.notebooklm.library_name,
        max_words_per_query=project.notebooklm.max_words_per_query,
        timeout_seconds=project.notebooklm.timeout_seconds,
    )
    spec = NotebookSectionSpec(
        id=section_id, title=title, focus=focus,
        priority_sources=[s.strip() for s in priority_sources.split(",") if s.strip()],
        target_chars=target_chars,
    )
    console.print(f"[dim]Asking NotebookLM for section {section_id} (~{target_chars} chars)…[/dim]")
    draft = write_section(adapter, spec)
    out = project.state_dir() / f"notebook_section_{section_id.replace('.', '_')}.md"
    out.write_text(
        f"# {section_id} {title}\n\n{draft.body}\n\n## Bibliografia wykorzystana\n\n"
        + "\n\n".join(draft.bibliography),
        encoding="utf-8",
    )
    console.print(f"[green]✓[/green] {len(draft.body)} chars, {len(draft.bibliography)} biblio entries")
    if draft.truncated:
        console.print("[yellow]⚠️ Response may be truncated — re-run with smaller target_chars.[/yellow]")
    console.print(f"Saved: [cyan]{out}[/cyan]")


@app.command("audit-citations")
def audit_citations_cmd(project_dir: Path) -> None:
    """STEP 2 — Haiku swarm citation audit (wraps Eryk's thesis-citation-audit skill)."""
    from thesis_generator.citation_audit import run_citation_audit

    project = ThesisProject.load(project_dir)
    console.print("[dim]Dispatching Haiku swarm (one sub-agent per subsection)…[/dim]")
    audit = run_citation_audit(project, max_parallel=project.pipeline.max_parallel_haiku)
    console.print(f"[bold]Total citations checked:[/bold] {audit.total_citations}")
    console.print(f"[bold]Problems flagged:[/bold] {audit.total_problems}")
    for s in audit.sections:
        prob = s.total - s.ok_count
        status = "[green]✓[/green]" if prob == 0 else f"[yellow]{prob} issues[/yellow]"
        console.print(f"  {s.section_id}: {s.ok_count}/{s.total} OK {status}")
    if audit.merged_report_path:
        console.print(f"Merged report: [cyan]{audit.merged_report_path}[/cyan]")


@app.command("orchestrate-fixes")
def orchestrate_fixes_cmd(project_dir: Path) -> None:
    """STEP 3 — Opus reads Haiku reports, builds structured fix list (analyst-then-builder pattern)."""
    from thesis_generator.citation_audit import orchestrate_fixes, run_citation_audit

    project = ThesisProject.load(project_dir)
    console.print("[dim]Re-loading audit + running Opus orchestrator…[/dim]")
    audit = run_citation_audit(project, max_parallel=project.pipeline.max_parallel_haiku)
    result = orchestrate_fixes(project, audit)
    console.print(f"[bold]Total fixes proposed:[/bold] {len(result.fixes)}")
    for section_id, fixes in result.by_section().items():
        console.print(f"  {section_id}: {len(fixes)} fix(es)")
    fix_json = project.state_dir() / "citation_fixes.json"
    console.print(f"Fix list saved: [cyan]{fix_json}[/cyan]")
    if result.summary:
        console.print(f"\n[bold]Executive summary:[/bold]\n{result.summary[:500]}…")


@app.command("notebook-pipeline")
def notebook_pipeline_cmd(
    project_dir: Path,
    skip: str = typer.Option("", help="Comma-separated steps to skip: write,audit,orchestrate,correct"),
) -> None:
    """Run all 4 steps end-to-end (NotebookLM-first pipeline)."""
    from thesis_generator.pipeline_notebook import run_notebook_pipeline

    project = ThesisProject.load(project_dir)
    skipped = tuple(s.strip() for s in skip.split(",") if s.strip())
    console.print(f"[dim]Running 4-step pipeline, skipping: {skipped or 'nothing'}[/dim]")
    result = run_notebook_pipeline(project, section_specs=None, skip_steps=skipped)

    console.print(f"\n[bold]Step 1 — drafts written:[/bold] {len(result.drafts)}")
    if result.audit:
        console.print(f"[bold]Step 2 — citations checked:[/bold] {result.audit.total_citations} "
                      f"({result.audit.total_problems} problems)")
    if result.fixes:
        console.print(f"[bold]Step 3 — fixes proposed:[/bold] {len(result.fixes.fixes)}")
    if result.corrections_by_section:
        console.print(f"[bold]Step 4 — sections corrected:[/bold] {len(result.corrections_by_section)}")
    for n in result.notes:
        console.print(f"[dim]NOTE: {n}[/dim]")


@app.command("visuals")
def visuals_cmd(
    project_dir: Path,
    output: Path = typer.Option(None, help="Write to this path instead of mutating draft.docx"),
) -> None:
    """Scan draft for visual markers and insert tables/charts/images + Spis tabel/rysunków."""
    from thesis_generator.visuals import process_visuals

    project = ThesisProject.load(project_dir)
    result = process_visuals(project, output_path=output)
    console.print(f"[bold]Inserted:[/bold] {result.inserted} visuals")
    if result.skipped:
        console.print(f"[yellow]Skipped:[/yellow] {len(result.skipped)}")
        for title, reason in result.skipped:
            console.print(f"  • [yellow]{title}[/yellow] — {reason}")
    if result.registry:
        console.print(
            f"[dim]Tables: {len(result.registry.tables)} | "
            f"Figures: {len(result.registry.figures)} | "
            f"Suggested (need spec): {len(result.registry.suggestions)}[/dim]"
        )
    for n in result.notes:
        console.print(f"  [dim]{n}[/dim]")
    console.print(f"Report: [cyan]{project.reports_dir() / 'visuals.md'}[/cyan]")
    if result.output_path:
        console.print(f"Output: [green]{result.output_path}[/green]")


@app.command("run")
def run_cmd(
    project_dir: Path,
    skip: str = typer.Option("", help="Comma-separated phases to skip"),
) -> None:
    """Run the full pipeline end-to-end."""
    from thesis_generator.pipeline import ALL_PHASES, run_pipeline

    project = ThesisProject.load(project_dir)
    skipped = {s.strip() for s in skip.split(",") if s.strip()}
    phases = tuple(p for p in ALL_PHASES if p not in skipped)
    result = run_pipeline(project, phases=phases)
    console.print(f"[bold]Phases completed:[/bold] {result.phases_completed}")
    if result.failed_phases:
        console.print(f"[red]Failed:[/red] {result.failed_phases}")
    if result.final_grade is not None:
        console.print(f"[bold]Reviewer grade:[/bold] {result.final_grade}")


@app.command("notebooklm")
def notebooklm_cmd(
    action: str = typer.Argument(..., help="status | auth | list"),
    project_dir: Path | None = typer.Option(None),
) -> None:
    """Manage the NotebookLM verification engine."""
    from thesis_generator.notebooklm import NotebookLMAdapter

    if project_dir:
        project = ThesisProject.load(project_dir)
        if project.notebooklm is None:
            console.print("[red]No notebooklm block in thesis.yaml[/red]")
            raise typer.Exit(1)
        url = project.notebooklm.library_url
        name = project.notebooklm.library_name
    else:
        url = ""
        name = ""

    adapter = NotebookLMAdapter(library_url=url, library_name=name) if url else None

    if action == "status":
        if not adapter:
            console.print("Pass --project-dir to check a configured library's auth status.")
            raise typer.Exit(1)
        ok = adapter.check_auth()
        console.print(f"NotebookLM auth: {'✅ authenticated' if ok else '❌ NOT authenticated'}")
        if not ok:
            console.print("[yellow]Run: python ~/.claude/skills/notebooklm/scripts/run.py auth_manager.py setup[/yellow]")
    elif action == "auth":
        import subprocess, sys

        skill = Path.home() / ".claude" / "skills" / "notebooklm"
        subprocess.run([sys.executable, "scripts/run.py", "auth_manager.py", "setup"], cwd=str(skill))
    elif action == "list":
        if not adapter:
            console.print("Pass --project-dir.")
            raise typer.Exit(1)
        for lib in adapter.list_libraries():
            console.print(lib)
    else:
        console.print(f"Unknown action: {action}")
        raise typer.Exit(1)


@app.command("version")
def version_cmd() -> None:
    console.print(f"thesis-generator {__version__}")


def _starter_yaml(title: str, author: str) -> str:
    """Generate a starter thesis.yaml with comments pointing at every knob."""
    return f"""# thesis-generator project config
# Edit every `<...>` placeholder before running the pipeline.

title: "{title}"
author: "{author}"
school: ""
promotor: ""
language: pl                 # pl | en
citation_style: apa7         # apa7 | vancouver (planned) | chicago (planned)

# All inputs/* paths follow the conventional layout created by `tg init`.
# Only the two REQUIRED paths must be set explicitly; the rest auto-discover
# from inputs/research_data/, inputs/school/, etc. when those folders exist.
inputs:
  draft: inputs/draft.docx                       # REQUIRED — your existing partial draft
  sources_dir: inputs/sources                    # REQUIRED — folder with academic PDFs you cite
  research_data_dir: inputs/research_data        # OPTIONAL — auto-discovers surveys/, interviews/, observations/, existing/
  school_dir: inputs/school                      # OPTIONAL — auto-discovers regulation.docx, brief.pdf
  # Override only if you put files outside the conventional layout:
  # raw_data: [inputs/survey_responses.xlsx]
  # interviews_dir: /some/other/path
  # regulation: inputs/school/Zarządzenie_117.docx
  # school_brief: inputs/school/system_prompt.pdf

# NotebookLM library URL — REQUIRED for Ring C verification.
# Create a notebook at https://notebooklm.google.com and upload the same PDFs
# that live in inputs/sources/ (the library must be a strict superset).
notebooklm:
  library_url: "https://notebooklm.google.com/notebook/<UUID>"
  library_name: "My thesis library"
  max_words_per_query: 400          # NotebookLM truncates longer queries silently
  parallel_queries: 3
  timeout_seconds: 300

models:
  writer: claude-opus-4-7
  auditor: claude-opus-4-7
  reviewer: claude-opus-4-7
  rewriter: claude-sonnet-4-6
  verifier_per_source: claude-haiku-4-5-20251001

humanization:
  remove_em_dashes: true            # ` — ` → `, ` (en-dashes preserved for page ranges)
  forbidden_words:
    - triangulacja
    - rygorystycznie
    - holistycznie
    - synergiczny
    - paradygmat
  sentence_length_variance: true
  chevron_quotes: true              # match draft's existing quote style
  fake_access_dates: false          # set true to randomize netografia access dates
  fake_date_range: ["2024-03-01", "2025-06-30"]

output:
  dir: output
  docx_name: "Praca.docx"
  pdf: true                         # requires pywin32 on Windows or libreoffice elsewhere
  versioned: true                   # keep Praca_v1.docx … Praca_vN.docx in output/_versions/

pipeline:
  section_by_section: true
  verification_rings: ["A", "B", "C"]
  re_verify_on_edit: true
  reviewer_passes: 1                # 3 for paranoid mode
  max_parallel_haiku: 12
  fail_on_grade_below: 4.0          # pipeline.ship refuses below this grade
"""


# ============================================================================
# README templates for the scaffolded inputs/ tree
# ============================================================================


def _readme_inputs() -> str:
    return """# inputs/

All files you provide to the thesis generator live here. The tool **never**
modifies anything in this folder — outputs go to `output/`, reports go to
`_reports/`, intermediate state goes to `_state/`.

## Folder map

| Folder | What goes here |
|---|---|
| [`draft.docx`](draft.docx) | Your existing partial thesis (even if just the intro) |
| [`sources/`](sources/) | Academic literature PDFs you cite — books, papers, reports |
| [`research_data/`](research_data/) | YOUR research data — surveys, interviews, observations, secondary data |
| [`school/`](school/) | Documents from the university — formatting regulation, templates, promotor's brief |
| [`notes/`](notes/) | Anything else — guidelines, correspondence, ideas |

Drop files where they belong; the tool auto-discovers the conventional layout
described in [docs/CONFIG.md](../../docs/CONFIG.md). You don't need to list
every file in `thesis.yaml` — only the bits you put in non-conventional places.
"""


def _readme_sources() -> str:
    return """# inputs/sources/

**Drop here:** PDFs of every academic source you cite in the thesis.

The Ring B verifier (`tg verify --rings B`) opens each cited PDF via
`pdfplumber` and checks whether the paraphrase in your draft actually matches
what the source says. So every `(Author, year)` you cite needs a matching PDF
here.

## Naming convention (recommended)

`Author_Year_ShortTitle.pdf` — the inventory step uses the filename as a
heuristic to map `(Author, year)` citations to PDFs.

Examples:
- `Hoch_2002_Product_Experience.pdf`
- `Trusov_et_al_2009_WoM_Effects.pdf`
- `Batorski_2006_Badania_internetowe.pdf`

## Subfolders are OK

Recurse-discovered. You can organize however you want, e.g.:
- `primary_literature/` — peer-reviewed papers + books
- `reports/` — industry, government, NGO reports
- `netografia/` — web sources saved as PDF

## NotebookLM library

Whatever lives here should also live in your NotebookLM library (configured in
`thesis.yaml` under `notebooklm.library_url`). Ring C verification compares
NotebookLM's source-grounded answer to your draft's paraphrase.

If a source isn't in the NotebookLM library, you'll get `UNKNOWN` verdicts
on it. Upload it via the NotebookLM web UI.
"""


def _readme_research_data() -> str:
    return """# inputs/research_data/

**Drop here:** YOUR research data — what you collected or compiled yourself.

This is distinct from `sources/` (which is what you CITE). The data audit
phase (`tg audit-data`) recomputes every numeric claim in your draft against
files here, catching off-by-N errors and silently-rounded percentages.

## Subfolders

| Folder | What goes here |
|---|---|
| [`surveys/`](surveys/) | Quantitative — CAWI/PAPI/CATI questionnaires as `.xlsx`, `.csv`, `.json` |
| [`interviews/`](interviews/) | Qualitative — IDI/FGI transcripts as `.pdf`, `.docx`, `.md` |
| [`observations/`](observations/) | Ethnography, field notes, diaries |
| [`existing/`](existing/) | Secondary data — GUS, GUS BDL, industry reports as `.xlsx` |

The data audit reads `.xlsx`, `.xls`, `.csv`, `.json` from `surveys/` and
`existing/`. Other formats are inventoried but not numerically audited.
"""


def _readme_research_surveys() -> str:
    return """# inputs/research_data/surveys/

**Drop here:** Survey response files — `.xlsx`, `.csv`, `.json`.

The data audit reads numeric cells from every file here and cross-checks them
against numeric claims in your draft (e.g. "65% kolejnych headsetów" → does
the cell really say 65% or is it 34%?).

## Typical contents

- Raw export from Google Forms / Survio / Webankieta / Microsoft Forms
- An analysis workbook with H1/H2/H3 sheets, stats (M, SD, χ², t, p, d)
- Pivoted summary tables
"""


def _readme_research_interviews() -> str:
    return """# inputs/research_data/interviews/

**Drop here:** Interview transcripts — `.pdf`, `.docx`, `.md`.

The inventory step indexes these (they show up in the source map). The
reviewer reads them when grading the methodology section.

## Naming convention (recommended)

`InterviewID_Pseudonym_Date.pdf` — e.g. `IDI01_Anna_2026-03-15.pdf`,
`FGI02_VR-users_2026-04-02.pdf`. Pseudonyms only, per ethics norms.
"""


def _readme_research_observations() -> str:
    return """# inputs/research_data/observations/

**Drop here:** Field notes, ethnographic logs, diaries, observation protocols.

Any format — `.md` is easiest for the reviewer to read, but `.pdf` and `.docx`
are accepted too.
"""


def _readme_research_existing() -> str:
    return """# inputs/research_data/existing/

**Drop here:** Secondary data you re-use — public datasets, industry reports
in tabular form, GUS / Eurostat exports.

`.xlsx` and `.csv` here are picked up by the data audit alongside your primary
survey data, so you can cite "GUS 2024 — 67% of population uses internet" and
the audit will verify the number matches the cell you exported.
"""


def _readme_visuals() -> str:
    return """# inputs/visuals/

**Drop here:** ready images for figures (ILUSTRACJA markers) — photos,
diagrams, screenshots, schemas. Anything that should appear in the thesis
as-is, without being generated from data.

Charts produced from your survey/research data are NOT here — they are
auto-generated by `tg visuals` into `output/_charts/` based on `[WYKRES ...]`
markers with a `[Dane: ...]` spec.

## Naming convention

Use a short, descriptive slug — the writer (or you) will reference these
filenames in markers:

    [ILUSTRACJA 1: Sensorama (1962)]
    [Źródło: Wikimedia Commons, domena publiczna]
    [Plik: visuals/sensorama.jpg]
    [Szerokość: 11cm]

If the marker says only `[Plik: sensorama.jpg]`, the pipeline auto-searches
this folder.

## Accepted formats

`.png`, `.jpg`, `.jpeg`, `.bmp`, `.gif`. SVG is NOT supported by python-docx;
convert to PNG first.

## Subfolders are OK

Organize however you want; the writer just references the relative path.
"""


def _readme_school() -> str:
    return """# inputs/school/

**Drop here:** Documents from the university.

| Filename (conventional) | What it is |
|---|---|
| `regulation.docx` | Formatting regulation — *Zarządzenie ... załącznik 3*. Drives style audit. |
| `brief.pdf` | Promotor's task brief / system prompt for the thesis. Drives writer scope. |
| `templates/` | Subfolder for title pages, declarations, footer templates |

If your university uses different filenames, set `inputs.regulation` and
`inputs.school_brief` explicitly in `thesis.yaml`.

## Why this matters

- **regulation.docx**: the independent reviewer reads this when grading. Without
  it, formatting violations won't be flagged.
- **brief.pdf**: the writer uses this as system-prompt context so generated
  sections match what your promotor asked for, not a generic thesis template.
"""


def _readme_notes() -> str:
    return """# inputs/notes/

**Drop here:** Anything else that doesn't fit elsewhere — promotor's email
guidelines, your own outline notes, helpful articles you don't cite formally.

The writer is aware of files here via `inputs.additional` in `thesis.yaml`
when you list them; otherwise this is a free-form scratch folder the tool
ignores.
"""


if __name__ == "__main__":
    app()
