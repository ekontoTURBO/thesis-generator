"""Audit-only pipeline — verify an existing thesis without writing or mutating it.

For the second use case: someone already has a finished thesis (written however
— by hand, by ChatGPT, by this tool, by a co-author) and just wants the most
thorough quality report we can produce. No generation, no mutation, no
corrections applied — just diagnostics.

Every phase here is read-only or works on a temp copy. The draft.docx in the
project's inputs/ folder is never touched. All outputs land in _reports/.

Final artifact: `_reports/THESIS_AUDIT.md` — one consolidated report covering
citations, bibliography, raw data, formatting, reviewer grade, visual register,
and a prioritized fix list synthesized by Opus.
"""

from __future__ import annotations

import shutil
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from thesis_generator.citation_audit import (
    CitationAuditResult,
    OpusOrchestratorResult,
    orchestrate_fixes,
    run_citation_audit,
)
from thesis_generator.config import ThesisProject
from thesis_generator.docx_ops.humanize import HumanizationStats, humanize_docx
from thesis_generator.env_check import EnvCheckResult, check_env
from thesis_generator.inventory import Inventory, build_inventory
from thesis_generator.numbers_audit import (
    HypothesisConsistencyResult,
    NumericRecomputeResult,
    check_hypothesis_consistency,
    run_recompute,
)
from thesis_generator.review.independent import ReviewerResult, run_independent_review
from thesis_generator.verify.data_audit import DataAuditResult, run_data_audit
from thesis_generator.verify.internal import RingAResult, run_ring_a
from thesis_generator.visuals.markers import VisualKind, parse_markers


@dataclass(slots=True)
class VisualsCounts:
    tables: int = 0
    charts: int = 0
    images: int = 0
    suggested: int = 0
    total_existing_tables: int = 0  # raw python-docx tables in the file


@dataclass(slots=True)
class AuditResult:
    env: EnvCheckResult | None = None
    inventory: Inventory | None = None
    ring_a: RingAResult | None = None
    citation_audit: CitationAuditResult | None = None
    fixes: OpusOrchestratorResult | None = None
    data_audit: DataAuditResult | None = None
    numbers_recompute: NumericRecomputeResult | None = None
    hypothesis_consistency: HypothesisConsistencyResult | None = None
    reviewer: ReviewerResult | None = None
    visuals: VisualsCounts | None = None
    humanize: HumanizationStats | None = None
    consolidated_report: Path | None = None
    skipped: list[tuple[str, str]] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


def run_audit(
    project: ThesisProject,
    *,
    deep: bool = True,
    skip: tuple[str, ...] = (),
) -> AuditResult:
    """Run every read-only audit phase. `deep=False` skips the expensive LLM steps.

    Phase keys (for `skip=`):
      env, inventory, ring_a, citation_audit, orchestrate, data, reviewer,
      visuals, humanize
    """
    result = AuditResult()

    # 0. env_check
    if "env" not in skip:
        result.env = check_env(project)

    # 1. inventory (needed by ring_a / data_audit / citation_audit)
    if "inventory" not in skip:
        try:
            result.inventory = build_inventory(project)
        except Exception as e:
            result.skipped.append(("inventory", str(e)))

    # 2. Ring A — internal regex/python-docx (read-only)
    if "ring_a" not in skip:
        try:
            result.ring_a = run_ring_a(project)
        except Exception as e:
            result.skipped.append(("ring_a", str(e)))

    # 3. Visuals register — read-only count of TABELA/WYKRES/ILUSTRACJA markers + raw tables
    if "visuals" not in skip:
        try:
            result.visuals = _count_visuals(project)
        except Exception as e:
            result.skipped.append(("visuals", str(e)))

    # 4. Humanize REPORT (operates on a temp copy — never touches original)
    if "humanize" not in skip:
        try:
            result.humanize = _humanize_dry_run(project)
        except Exception as e:
            result.skipped.append(("humanize", str(e)))

    # 5. Data audit — read-only numeric cross-check
    if "data" not in skip:
        try:
            result.data_audit = run_data_audit(project)
        except Exception as e:
            result.skipped.append(("data_audit", str(e)))

    # 6. Numbers recompute (Opus, slow) — independent xlsx recomputation
    if deep and "numbers_recompute" not in skip:
        try:
            result.numbers_recompute = run_recompute(project)
        except Exception as e:
            result.skipped.append(("numbers_recompute", str(e)))

    # 6b. Hypothesis consistency — only if recompute succeeded
    if (
        deep
        and "hypothesis_consistency" not in skip
        and result.numbers_recompute is not None
        and result.numbers_recompute.diffs
    ):
        try:
            result.hypothesis_consistency = check_hypothesis_consistency(
                project, result.numbers_recompute
            )
        except Exception as e:
            result.skipped.append(("hypothesis_consistency", str(e)))

    # 7. Citation audit (Haiku swarm) — slow, deep mode only
    if deep and "citation_audit" not in skip:
        try:
            result.citation_audit = run_citation_audit(project)
        except Exception as e:
            result.skipped.append(("citation_audit", str(e)))

    # 7. Opus orchestrator — only useful if audit ran
    if deep and "orchestrate" not in skip and result.citation_audit is not None:
        try:
            result.fixes = orchestrate_fixes(project, result.citation_audit)
        except Exception as e:
            result.skipped.append(("orchestrate", str(e)))

    # 8. Independent reviewer — read-only Opus single call
    if deep and "reviewer" not in skip:
        try:
            result.reviewer = run_independent_review(project)
        except Exception as e:
            result.skipped.append(("reviewer", str(e)))

    # 9. Consolidate
    out = project.reports_dir() / "THESIS_AUDIT.md"
    out.write_text(_render_consolidated_report(result, project), encoding="utf-8")
    result.consolidated_report = out

    return result


# ============================================================================
# Read-only helpers
# ============================================================================


def _count_visuals(project: ThesisProject) -> VisualsCounts:
    """Count visual markers + existing tables without mutating anything."""
    draft = project.resolve_input(project.inputs.draft)
    doc = Document(str(draft))
    counts = VisualsCounts(total_existing_tables=len(doc.tables))
    text = "\n".join(p.text for p in doc.paragraphs)
    for m in parse_markers(text):
        if m.kind == VisualKind.TABLE:
            counts.tables += 1
        elif m.kind == VisualKind.CHART:
            counts.charts += 1
        elif m.kind == VisualKind.IMAGE:
            counts.images += 1
        else:
            counts.suggested += 1
    return counts


def _humanize_dry_run(project: ThesisProject) -> HumanizationStats:
    """Run humanize on a temp COPY of the draft so the original is untouched."""
    draft = project.resolve_input(project.inputs.draft)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        tmp = Path(tf.name)
    shutil.copy(draft, tmp)
    try:
        return humanize_docx(tmp, project.humanization)
    finally:
        try:
            tmp.unlink(missing_ok=True)
        except Exception:
            pass


# ============================================================================
# Consolidated report renderer
# ============================================================================


def _render_consolidated_report(result: AuditResult, project: ThesisProject) -> str:
    lines: list[str] = []
    lines.append(f"# THESIS AUDIT — {project.title}")
    lines.append(f"\n**Author:** {project.author}")
    lines.append(f"**Project:** `{project.project_dir}`")
    lines.append(f"**Draft:** `{project.resolve_input(project.inputs.draft)}`\n")

    # Executive summary — the headline numbers + grade
    lines.append("## Executive summary\n")
    if result.reviewer and result.reviewer.grade is not None:
        lines.append(f"- **Independent reviewer grade:** {result.reviewer.grade}")
    if result.ring_a:
        lines.append(f"- **Ring A status:** {'✅ PASS' if result.ring_a.passed else '❌ FAIL'}")
        lines.append(f"  - {len(result.ring_a.citations_in_text)} citations, {len(result.ring_a.bib_entries)} bib entries")
        lines.append(f"  - missing from bibliography: {len(result.ring_a.missing_in_bib)}")
        lines.append(f"  - orphaned bibliography entries: {len(result.ring_a.orphaned_in_bib)}")
        lines.append(f"  - em-dashes (AI tell): {result.ring_a.em_dash_count}")
    if result.citation_audit:
        lines.append(
            f"- **Citation audit (Haiku swarm):** "
            f"{result.citation_audit.total_problems} problems / "
            f"{result.citation_audit.total_citations} citations checked"
        )
    if result.fixes and result.fixes.fixes:
        lines.append(f"- **Suggested fixes (Opus):** {len(result.fixes.fixes)}")
    if result.data_audit:
        lines.append(
            f"- **Data audit (heuristic):** {len(result.data_audit.likely_off_by)} likely mismatches "
            f"out of {len(result.data_audit.claims)} numeric claims"
        )
    if result.numbers_recompute:
        nr = result.numbers_recompute
        lines.append(
            f"- **Numbers recompute (Opus, independent xlsx):** {nr.ok_count}/{nr.total} OK, "
            f"{nr.mismatches} mismatches, {nr.unverifiable} unverifiable"
        )
    if result.hypothesis_consistency:
        hc = result.hypothesis_consistency
        lines.append(
            f"- **Hypothesis consistency:** {hc.supported}/{len(hc.verdicts)} supported, "
            f"{hc.overinterpretations} over-interpreted"
        )
    if result.visuals:
        v = result.visuals
        lines.append(
            f"- **Visuals:** {v.total_existing_tables} tables, "
            f"{v.tables} TABELA markers, {v.charts} WYKRES, {v.images} ILUSTRACJA, "
            f"{v.suggested} suggested"
        )
    if result.humanize:
        lines.append(
            f"- **Humanize check (dry-run):** {result.humanize.em_dashes_replaced} em-dashes flagged, "
            f"{len(result.humanize.forbidden_words_found)} forbidden buzzwords found"
        )

    # Opus thesis-wide summary if available
    if result.fixes and result.fixes.summary:
        lines.append(f"\n### Opus orchestrator synthesis\n\n{result.fixes.summary}")

    # Per-area sections
    if result.ring_a:
        lines.append("\n## Ring A — Internal audit\n")
        lines.append(result.ring_a.to_report())

    if result.citation_audit:
        lines.append("\n## Citation audit — per-section\n")
        lines.append("| Section | OK | ⚠️ STRONA | ❌ TREŚĆ | ❓ BRAK | 🔴 BIBLIO |")
        lines.append("|---|---|---|---|---|---|")
        for s in result.citation_audit.sections:
            lines.append(
                f"| {s.section_id} | {s.ok_count} | {s.page_issues} | {s.content_issues} | "
                f"{s.missing_source} | {s.biblio_issues} |"
            )
        if result.citation_audit.merged_report_path:
            lines.append(
                f"\n_Full per-citation detail in `{result.citation_audit.merged_report_path}`_"
            )

    if result.fixes and result.fixes.fixes:
        lines.append("\n## Suggested fixes (Opus)\n")
        lines.append("| Section | Paragraph | Flag | Cited as | Hint |")
        lines.append("|---|---|---|---|---|")
        for f in result.fixes.fixes[:50]:  # cap to keep report readable
            hint = f.correction_hint.replace("|", "\\|").replace("\n", " ")[:90]
            lines.append(f"| {f.section_id} | {f.paragraph_tag} | {f.flag} | {f.cited_as} | {hint} |")
        if len(result.fixes.fixes) > 50:
            lines.append(f"\n_… {len(result.fixes.fixes) - 50} more fixes in `_state/citation_fixes.json`_")

    if result.data_audit and result.data_audit.likely_off_by:
        lines.append("\n## Data audit — likely numeric mismatches (heuristic)\n")
        for c, ev, reason in result.data_audit.likely_off_by[:30]:
            lines.append(f"- `{c.raw_text}` (paragraph {c.paragraph_idx}) — Excel `{ev.sheet}!{ev.cell}` = {ev.value}. {reason}")

    if result.numbers_recompute and result.numbers_recompute.diffs:
        nr = result.numbers_recompute
        mismatches = [d for d in nr.diffs if d.status == "MISMATCH"]
        if mismatches:
            lines.append("\n## Numbers recompute — confirmed MISMATCHES (Opus independent calculation)\n")
            lines.append("| Paragraph | Claimed | Recomputed | Δ | Note |")
            lines.append("|---|---|---|---|---|")
            for d in mismatches[:50]:
                note = d.note.replace("|", "\\|").replace("\n", " ")[:120]
                lines.append(f"| {d.paragraph_tag} | {d.claimed} | {d.recomputed} | {d.delta} | {note} |")
        if nr.workbook_path:
            lines.append(f"\n_Full parallel workbook:_ `{nr.workbook_path}`")

    if result.hypothesis_consistency and result.hypothesis_consistency.verdicts:
        lines.append("\n## Hypothesis ↔ conclusion consistency\n")
        for v in result.hypothesis_consistency.verdicts:
            icon = {"SUPPORTED": "✅", "NOT_SUPPORTED": "❌", "PARTIALLY": "⚠️",
                    "OVER_INTERPRETED": "🚨", "INSUFFICIENT_DATA": "❓"}.get(v.judged_verdict, "•")
            lines.append(f"\n### {icon} {v.hypothesis_id} — {v.judged_verdict}")
            lines.append(f"_Declaration:_ {v.declaration}")
            if v.claimed_verdict:
                lines.append(f"_Author claims:_ {v.claimed_verdict}")
            lines.append(f"\n{v.reasoning}")

    if result.reviewer:
        lines.append("\n## Independent reviewer report\n")
        lines.append(result.reviewer.to_report())

    if result.humanize and result.humanize.forbidden_words_found:
        lines.append("\n## AI tells / forbidden buzzwords\n")
        for w in result.humanize.forbidden_words_found:
            lines.append(f"- {w}")

    if result.skipped:
        lines.append("\n## Phases skipped\n")
        for phase, reason in result.skipped:
            lines.append(f"- **{phase}**: {reason}")

    lines.append("\n---\n")
    lines.append(
        "_This is an AUDIT report — no part of the draft was modified. "
        "All findings are claims to verify, not facts to silently apply._"
    )
    return "\n".join(lines)
