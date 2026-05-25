"""Apply edits to a .docx from a JSON decision file.

This is the JSON-decision-file pattern from brief section 5: subagents that
need to make many small edits NEVER touch the docx directly. They emit a JSON
file like:

    {
      "edits": [
        {"para": 42, "old": "old phrase", "new": "new phrase"},
        ...
      ]
    }

Then the main session calls this function to apply them. This avoids the
string-escape disaster (gotcha #8): when subagents write JSON manually with
f-strings, ASCII quotes in values blow up parsing. Forcing `json.dump()` on
the subagent side + dedicated applier on this side eliminates the bug class.

Ported from `apply_lang_v13.py` — preserves the "track changes" visualization:
old text strikethrough+gray, new text green.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass(slots=True)
class EditApplyResult:
    applied: int = 0
    skipped: list[tuple[int, str, str]] = field(default_factory=list)
    """list of (paragraph_idx, reason, old_text_preview)"""
    numeric_or_citation_drift: list[str] = field(default_factory=list)

    def render(self) -> str:
        lines = [f"# Edit application — {self.applied} applied, {len(self.skipped)} skipped\n"]
        if self.skipped:
            lines.append("## Skipped\n")
            for para, reason, old in self.skipped:
                lines.append(f"- p.{para} `{reason}` — `{old[:60]}`")
        if self.numeric_or_citation_drift:
            lines.append("\n## ⚠️ Edits that may have changed numbers/citations (review)\n")
            for d in self.numeric_or_citation_drift:
                lines.append(f"- {d}")
        return "\n".join(lines)


_GRAY = RGBColor(0x80, 0x80, 0x80)
_GREEN = RGBColor(0x1F, 0x7A, 0x1F)
_BLACK = RGBColor(0, 0, 0)


def _add_run(paragraph, text: str, color: RGBColor, *, strike: bool = False, font: str = "Arial", size: int = 12) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.strike = strike
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), font)


def apply_edits_from_json(
    docx_path: Path,
    edits_json: Path,
    output_path: Path | None = None,
    *,
    track_changes_visual: bool = True,
    safety_check_numbers: bool = True,
) -> EditApplyResult:
    """Apply edits described in `edits_json` to `docx_path`."""
    data = json.loads(edits_json.read_text(encoding="utf-8"))
    edits = data["edits"]

    # Group edits per paragraph
    by_para: dict[int, list[dict]] = {}
    for e in edits:
        by_para.setdefault(e["para"], []).append(e)

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs
    result = EditApplyResult()

    for pi, para_edits in sorted(by_para.items()):
        if pi >= len(paragraphs):
            result.skipped.append((pi, "paragraph index out of range", str(para_edits)[:80]))
            continue
        p = paragraphs[pi]
        text = p.text

        # Locate each old-string within the paragraph
        spans: list[tuple[int, int, dict]] = []
        for e in para_edits:
            old = e["old"]
            idx = text.find(old)
            if idx < 0:
                result.skipped.append((pi, "old string NOT found", old[:60]))
                continue
            spans.append((idx, idx + len(old), e))
        spans.sort()

        # Drop overlapping spans (keep first-come)
        clean: list[tuple[int, int, dict]] = []
        last_end = -1
        for s, en, e in spans:
            if s < last_end:
                result.skipped.append((pi, "overlap with prior edit", e["old"][:60]))
                continue
            clean.append((s, en, e))
            last_end = en

        if not clean:
            continue

        # Clear runs then rebuild with colored runs
        style = p.style
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        cur = 0
        for s, en, e in clean:
            _add_run(p, text[cur:s], _BLACK)
            if track_changes_visual:
                _add_run(p, text[s:en], _GRAY, strike=True)
                _add_run(p, e["new"], _GREEN)
            else:
                _add_run(p, e["new"], _BLACK)
            cur = en
            result.applied += 1
            if safety_check_numbers:
                _check_numeric_drift(e, result, pi)
        _add_run(p, text[cur:], _BLACK)
        p.style = style

    doc.save(str(output_path or docx_path))
    return result


def _check_numeric_drift(edit: dict, result: EditApplyResult, para_idx: int) -> None:
    """Flag any edit that changed the number or citation set inside it."""
    old_nums = sorted(re.findall(r"\d+[.,]?\d*", edit["old"]))
    new_nums = sorted(re.findall(r"\d+[.,]?\d*", edit["new"]))
    old_cites = sorted(re.findall(r"\([^)]*\d{4}[^)]*\)", edit["old"]))
    new_cites = sorted(re.findall(r"\([^)]*\d{4}[^)]*\)", edit["new"]))
    if old_nums != new_nums or old_cites != new_cites:
        result.numeric_or_citation_drift.append(
            f"p.{para_idx}: `{edit['old'][:50]}` → `{edit['new'][:50]}`"
        )
