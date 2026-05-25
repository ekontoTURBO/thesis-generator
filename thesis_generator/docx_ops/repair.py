"""Repair common Word-file disasters.

Currently handles the **Ctrl+F9 wrap disaster** (brief gotcha #5): user
accidentally selects all + presses Ctrl+F9, which wraps the entire body in
a `w:fldChar` field code. Word then renders the document as blank because
it interprets the body text as a field instruction.

Two-pass repair (ported from the proven session's `repair_v10.py` +
`unwrap_outer_field.py`):

  1. **Pass A — instrText → t**: walk every `<w:instrText>` element. If its
     content is NOT a real field instruction (TOC / PAGE / PAGEREF / HYPERLINK
     / SEQ / REF / STYLEREF / NUMPAGES), convert it to a visible `<w:t>` text
     run. This restores the corrupted body text.

  2. **Pass B — unwrap outer field shells**: walk `<w:fldChar>` pairs. Any
     depth-0 pair whose instruction is NOT a real field code is the leftover
     "frame" of the disaster — remove it while preserving everything inside.

Both passes preserve legitimate Word fields (TOC, PAGEREF, hyperlinks, etc.).
A backup copy is always written before any modification.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

import lxml.etree as ET
from docx import Document
from docx.oxml.ns import qn


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XMLSPACE_NS = "http://www.w3.org/XML/1998/namespace"
_WT = f"{{{_W_NS}}}t"
_WIT = f"{{{_W_NS}}}instrText"
_FLDCHARTYPE = f"{{{_W_NS}}}fldCharType"
_XMLSPACE = f"{{{_XMLSPACE_NS}}}space"

# Strings that count as legitimate Word field instructions and MUST be preserved.
_REAL_FIELD_RE = re.compile(
    r"^(TOC\b|PAGE\b|PAGEREF\b|HYPERLINK\b|SEQ\b|REF\b|STYLEREF\b|NUMPAGES\b)",
)


@dataclass(slots=True)
class RepairReport:
    diagnosed: list[str] = field(default_factory=list)
    fixes_applied: list[str] = field(default_factory=list)
    backup_path: Path | None = None
    needs_manual: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return not self.needs_manual

    def render(self) -> str:
        lines = ["# Word document repair\n"]
        if self.diagnosed:
            lines.append("## Diagnosed\n")
            lines.extend(f"- {d}" for d in self.diagnosed)
        if self.fixes_applied:
            lines.append("\n## Fixed\n")
            lines.extend(f"- ✅ {f}" for f in self.fixes_applied)
        if self.backup_path:
            lines.append(f"\n_Backup saved to_ `{self.backup_path}`")
        if self.needs_manual:
            lines.append("\n## ❌ Needs manual intervention\n")
            lines.extend(f"- {m}" for m in self.needs_manual)
        return "\n".join(lines)


def diagnose(path: Path) -> RepairReport:
    """Inspect the docx for known disaster signatures. Read-only."""
    report = RepairReport()
    doc = Document(str(path))
    body = doc.element.body

    instr_count = len(body.findall(f".//{qn('w:instrText')}"))
    text_count = len(body.findall(f".//{qn('w:t')}"))
    fldchar_begins = len(
        body.findall(f".//{qn('w:fldChar')}[@{qn('w:fldCharType')}='begin']")
    )

    report.diagnosed.append(
        f"w:instrText elements: {instr_count}, w:t elements: {text_count}, "
        f"w:fldChar begin elements: {fldchar_begins}"
    )
    if instr_count > 100 and text_count < instr_count // 4:
        report.diagnosed.append(
            "🚨 Pattern matches Ctrl+F9 wrap disaster: "
            "huge field-code count vs tiny actual-text count."
        )
        report.needs_manual.append(
            "Run `repair_docx(path)` to unwrap the outer field shell."
        )

    return report


def repair_docx(path: Path, output_path: Path | None = None) -> RepairReport:
    """Apply both repair passes (instrText→t, then unwrap outer shells).

    Always writes a `_corrupted_backup.docx` next to the input before modifying
    anything. If `output_path` is given, the repaired file lands there;
    otherwise the input is overwritten in place.
    """
    path = Path(path)
    report = diagnose(path)
    target = Path(output_path) if output_path else path

    # Always back up before touching anything.
    backup = path.with_name(path.stem + "_corrupted_backup" + path.suffix)
    if not backup.exists():
        shutil.copyfile(path, backup)
    report.backup_path = backup

    # If diagnose says nothing wrong, no-op (but we already saved the backup).
    if not report.needs_manual:
        return report

    # Load the raw zip; we operate on word/document.xml directly to bypass
    # python-docx's element-tracking layer (which can choke on the corruption).
    with zipfile.ZipFile(path, "r") as zin:
        items = {n: zin.read(n) for n in zin.namelist()}
    root = ET.fromstring(items["word/document.xml"])
    body = root.find(f"{{{_W_NS}}}body")

    # ---- Pass A: convert orphaned instrText back to visible text ----
    converted = kept = 0
    for el in root.iter(_WIT):
        s = (el.text or "").strip()
        if s == "" or _REAL_FIELD_RE.match(s):
            kept += 1
            continue
        el.tag = _WT
        el.set(_XMLSPACE, "preserve")
        converted += 1
    report.fixes_applied.append(
        f"Pass A: converted {converted} <w:instrText> back to <w:t> "
        f"(kept {kept} real field instructions)."
    )

    # ---- Pass B: walk fldChar frames, remove depth-0 fake shells ----
    stack: list[dict] = []
    frames: list[dict] = []
    for el in body.iter():
        ln = ET.QName(el).localname
        if ln == "fldChar":
            ft = el.get(_FLDCHARTYPE)
            if ft == "begin":
                fr = {"depth": len(stack), "instr": [], "elems": [el]}
                frames.append(fr)
                stack.append(fr)
            elif ft in ("separate", "end"):
                if stack:
                    stack[-1]["elems"].append(el)
                    if ft == "end":
                        stack.pop()
        elif ln == "instrText":
            if stack:
                stack[-1]["instr"].append(el.text or "")
                stack[-1]["elems"].append(el)

    removed_shells = 0
    for fr in frames:
        code = "".join(fr["instr"]).strip()
        if fr["depth"] == 0 and not _REAL_FIELD_RE.match(code):
            for e in fr["elems"]:
                par = e.getparent()
                if par is not None:
                    par.remove(e)
            removed_shells += 1
    report.fixes_applied.append(
        f"Pass B: removed {removed_shells} outer field shells "
        f"(left {len(frames) - removed_shells} legit fields intact)."
    )

    # Write back
    items["word/document.xml"] = ET.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in items.items():
            zout.writestr(n, data)

    # Final sanity check via python-docx
    try:
        d = Document(str(target))
        ne = sum(1 for p in d.paragraphs if p.text.strip())
        report.fixes_applied.append(
            f"Post-repair: {len(d.paragraphs)} paragraphs, {ne} non-empty, "
            f"{len(d.tables)} tables, {len(d.inline_shapes)} images."
        )
        report.needs_manual = []  # success — clear the manual flag
    except Exception as e:
        report.needs_manual.append(
            f"Post-repair sanity check failed: {e}. Restore from {backup}."
        )
    return report
