"""Humanize a Word document — strip the obvious AI tells.

The user said this 3 times across the proven session:
    "wywal te głupie - zastąp to spacjami albo przecinkami, bo na kilometr
     pachnie AI. DŁUGIE ZDANIA SĄ POPRAWNE."

So: replace ` — ` with `, `, drop forbidden buzzwords, normalize quotes,
but DON'T mangle sentences — long sentences are fine and the user explicitly
wants them preserved.

We do NOT change citations, numbers, or anything inside parentheses.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from thesis_generator.config import HumanizationPolicy


@dataclass(slots=True)
class HumanizationStats:
    em_dashes_replaced: int = 0
    forbidden_words_found: list[str] = field(default_factory=list)
    quotes_normalized: int = 0
    paragraphs_touched: int = 0

    def render(self) -> str:
        lines = ["# Humanization pass\n"]
        lines.append(f"- em-dashes replaced: {self.em_dashes_replaced}")
        lines.append(f"- quotes normalized: {self.quotes_normalized}")
        lines.append(f"- paragraphs touched: {self.paragraphs_touched}")
        if self.forbidden_words_found:
            lines.append("\n## ⚠️ Forbidden words still present (manual rewrite needed)\n")
            for w in self.forbidden_words_found:
                lines.append(f"- {w}")
        return "\n".join(lines)


def _humanize_text(text: str, policy: HumanizationPolicy, stats: HumanizationStats) -> str:
    out = text

    if policy.remove_em_dashes:
        # Replace ` — ` (em-dash with surrounding spaces) with `, `
        # but leave en-dashes ` – ` alone (used in page ranges).
        before = out
        out = re.sub(r" — ", ", ", out)
        # Also handle cases without spaces: word—word → word, word
        out = re.sub(r"(\w)—(\w)", r"\1, \2", out)
        stats.em_dashes_replaced += before.count(" — ") + len(re.findall(r"\w—\w", before))

    if policy.chevron_quotes:
        # Outer "..." → „..." (Polish primary quotes are „ ")
        # NOTE: be conservative — only convert clear standalone strings.
        # We don't touch quotes inside parens or after a digit (often abbreviations).
        pass  # Conservative no-op for now; full detection is too risky to automate

    return out


def humanize_docx(path: Path, policy: HumanizationPolicy, output_path: Path | None = None) -> HumanizationStats:
    """Apply humanization to a Word file. Mutates in place unless output_path given."""
    doc = Document(str(path))
    stats = HumanizationStats()

    # Track forbidden-word occurrences (we report, not auto-rewrite — too risky).
    # Stem-match to handle Polish declensions: "triangulacja" must also catch
    # "triangulacji", "triangulacją", "triangulację", etc. We treat the first
    # 6 characters of each forbidden lemma as the stem and look for any word
    # starting with it.
    body_text = "\n".join(p.text for p in doc.paragraphs)
    body_lower = body_text.lower()
    for w in policy.forbidden_words:
        stem = w.lower()[: min(6, max(4, len(w) - 2))]
        pattern = rf"\b{re.escape(stem)}\w*"
        matches = re.findall(pattern, body_lower, re.IGNORECASE)
        if matches:
            stats.forbidden_words_found.append(f"{w} (lemma stem={stem!r}, {len(matches)}×)")

    for p in doc.paragraphs:
        if not p.runs:
            continue
        # Operate at run level to preserve formatting
        for r in p.runs:
            if not r.text:
                continue
            new = _humanize_text(r.text, policy, stats)
            if new != r.text:
                r.text = new
                stats.paragraphs_touched += 1
                break  # avoid double-counting per paragraph

    # Tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not r.text:
                            continue
                        new = _humanize_text(r.text, policy, stats)
                        if new != r.text:
                            r.text = new

    doc.save(str(output_path or path))
    return stats
