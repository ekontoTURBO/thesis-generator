"""Extract numeric claims + hypothesis statements from a thesis draft.

We feed Opus a STRUCTURED list of claims, not the whole thesis text, so the
prompt stays compact and Opus can focus on what matters (every statistic,
every H1/H2/H3 declaration, every "wykazano że..." sentence).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


@dataclass(slots=True)
class NumericClaim:
    """A specific numeric assertion in the draft."""

    paragraph_idx: int
    sentence: str                        # the whole sentence containing the number
    numbers_found: list[str]             # raw matches like "65%", "M = 4.21", "χ² = 12.3"
    citation_nearby: str | None = None   # e.g. "(opracowanie własne, n=120)"

    def to_prompt_line(self) -> str:
        nums = ", ".join(self.numbers_found)
        return f"[P{self.paragraph_idx:04d}] {nums}  ←  {self.sentence}"


@dataclass(slots=True)
class HypothesisStatement:
    """One H1/H2/H3 declaration or its later confirmation/rejection statement."""

    kind: str                            # "DECLARATION" | "TEST_RESULT" | "CONCLUSION"
    hypothesis_id: str                   # "H1", "H2", "H1a", etc.
    paragraph_idx: int
    sentence: str
    verdict_keyword: str | None = None   # "potwierdzona", "odrzucona", "częściowo"

    def to_prompt_line(self) -> str:
        v = f" [{self.verdict_keyword}]" if self.verdict_keyword else ""
        return f"[P{self.paragraph_idx:04d}] {self.kind} {self.hypothesis_id}{v}: {self.sentence}"


# --- Numeric pattern matching ----------------------------------------------

# Statistics patterns the proven session used:
# M = 4.21, SD = 1.13, N = 105, n = 62, χ² = 12.3, t = 2.45, p = 0.003, d = 0.42,
# V = 0.21 (Craméra), r = 0.34, %≥4 = 65, p < 0.05, F(2, 100) = 4.2
_STAT_PATTERNS = [
    re.compile(r"\b[MN]\s*=\s*\d+(?:[.,]\d+)?", re.IGNORECASE),
    re.compile(r"\bn\s*=\s*\d+", re.IGNORECASE),
    re.compile(r"\bSD\s*=\s*\d+(?:[.,]\d+)?", re.IGNORECASE),
    re.compile(r"\bχ²\s*[<>=]+\s*\d+(?:[.,]\d+)?"),
    re.compile(r"\bchi[-\s]?square\s*[<>=]+\s*\d+(?:[.,]\d+)?", re.IGNORECASE),
    re.compile(r"\b[tFr]\s*[<>=]\s*-?\d+(?:[.,]\d+)?"),
    re.compile(r"\bp\s*[<>=]+\s*0?[.,]\d+"),
    re.compile(r"\bp\s*[<>=]+\s*\d+e-?\d+", re.IGNORECASE),
    re.compile(r"\bd\s*=\s*-?\d+(?:[.,]\d+)?"),
    re.compile(r"\bV\s*[(=].*?\d+(?:[.,]\d+)?", re.IGNORECASE),
    re.compile(r"\d+(?:[.,]\d+)?\s*%"),               # 65%, 12.3%
    re.compile(r"\d+\s*(?:respondent\w*|osób|os\.|uczestnik\w*|badany\w*)", re.IGNORECASE),
    re.compile(r"\d+\s*(?:z|na)\s+\d+\b"),            # "4 z 120"
]


def _split_sentences(text: str) -> list[str]:
    """Naive sentence split that respects Polish abbreviations (s., r., op.cit.)."""
    # Protect known abbreviations
    protected = re.sub(r"\b(s|r|op\.\s*cit|por|red|tłum|tj|np|cyt)\.", lambda m: m.group(0).replace(".", "###DOT###"), text)
    # Split on sentence-ending punctuation
    parts = re.split(r"(?<=[.!?])\s+(?=[A-ZŁŚŻŹĆ])", protected)
    return [p.replace("###DOT###", ".").strip() for p in parts if p.strip()]


def extract_numeric_claims(draft: Path) -> list[NumericClaim]:
    """Walk the draft, return every sentence containing at least one numeric statistic."""
    doc = Document(str(draft))
    out: list[NumericClaim] = []
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        for sentence in _split_sentences(p.text):
            matches: list[str] = []
            for pat in _STAT_PATTERNS:
                for m in pat.finditer(sentence):
                    s = m.group(0).strip()
                    if s not in matches:
                        matches.append(s)
            if not matches:
                continue
            # Skip pure page citations / years
            if all(re.fullmatch(r"19\d{2}|20\d{2}|\(?\d+%?\)?", x) for x in matches):
                # Only year-like or trivial single numbers — skip unless it has a % or stat letter
                if not any("%" in x or re.search(r"[a-zA-Z]", x) for x in matches):
                    continue
            out.append(NumericClaim(
                paragraph_idx=i,
                sentence=sentence[:400],
                numbers_found=matches,
            ))
    return out


# --- Hypothesis pattern matching -------------------------------------------

_H_DECLARATION = re.compile(
    r"\b(H\d[a-z]?)\b\s*[:.]?\s*(.+?)(?=(?:\.\s)|$)",
    re.IGNORECASE | re.UNICODE,
)
_H_VERDICT_KEYWORDS = [
    "potwierdzon", "odrzucon", "częściowo potwierdz", "nie znaleziono podstaw",
    "supported", "rejected", "partially supported", "not supported",
]


def extract_hypotheses(draft: Path) -> list[HypothesisStatement]:
    """Find every H1/H2/H3 mention with its verdict keyword if present."""
    doc = Document(str(draft))
    out: list[HypothesisStatement] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        for sentence in _split_sentences(text):
            for m in _H_DECLARATION.finditer(sentence):
                hid = m.group(1).upper()
                content = m.group(2)
                verdict = None
                lower = sentence.lower()
                kind = "DECLARATION"
                for kw in _H_VERDICT_KEYWORDS:
                    if kw in lower:
                        verdict = kw
                        kind = "TEST_RESULT"
                        break
                # Heuristic: if sentence starts with "Wyniki potwierdzają..." or similar → CONCLUSION
                if kind == "DECLARATION" and re.search(
                    r"\b(wyniki|dane|analiza|test\w*|results?)\b.*\b" + hid + r"\b",
                    lower,
                ):
                    kind = "CONCLUSION"
                out.append(HypothesisStatement(
                    kind=kind,
                    hypothesis_id=hid,
                    paragraph_idx=i,
                    sentence=sentence[:400],
                    verdict_keyword=verdict,
                ))
    return out
