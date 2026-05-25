"""Independent reviewer — zero-context Opus, prompted as strict professor.

The proven session ran this 3 times (after v5, v9, v12). Each time produced
a concrete prioritized fix list + numeric grade (2.0–5.0). The grade became
a fixed metric to track quality across versions.

Key design: the reviewer subagent gets ONLY the thesis text + (optionally)
the school's regulation. NEVER the session history, never the prior reviewer
reports, never the verification reports. This forces fresh-eyes critique.

Gotcha #12: reviewer output is a list of CLAIMS to verify, not facts to apply.
A follow-up `verify_reviewer_claims` pass cross-checks each claim against the
real PDFs and marks it TRUE / HALLUCINATION / PARTIAL / ALREADY FIXED.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from thesis_generator.config import ThesisProject
from thesis_generator.llm import ClaudeCLI


REVIEWER_SYSTEM_PROMPT = """Jesteś doświadczonym profesorem ekonomii / nauk społecznych,
recenzentem prac licencjackich. Surowy, ale sprawiedliwy. Oceniasz w skali 2-5
(2 = niedostateczna, 3 = dostateczna, 4 = dobra, 4.5 = dobra plus, 5 = bardzo dobra).

NIE MASZ żadnego kontekstu jak ta praca powstawała. Zobaczysz tylko jej tekst
plus (opcjonalnie) regulamin uczelni. Twoje zadanie:

1. Przeczytaj pracę krytycznie, jakby trafiła Ci do recenzji na ślepo.
2. Wypisz TOP-5 najpoważniejszych problemów (priorytet: PILNY / WYSOKI / ŚREDNI).
   Dla każdego problemu:
   - dokładny cytat z pracy
   - dlaczego to problem (1-2 zdania)
   - sugerowana naprawa (1 zdanie)
3. Wypisz 3 mocne strony.
4. Wystaw OCENĘ z uzasadnieniem (3-4 zdania).

Format odpowiedzi:

# Recenzja niezależna

## Ocena: X.X
<uzasadnienie>

## Top problemy
### 1. [PILNY] <tytuł>
**Cytat:** "<dosłownie>"
**Dlaczego problem:** <…>
**Naprawa:** <…>

### 2. [WYSOKI] ...

## Mocne strony
- ...
- ...
- ...
"""


@dataclass(slots=True)
class ReviewerProblem:
    priority: str  # PILNY | WYSOKI | ŚREDNI
    title: str
    quote: str
    why: str
    fix: str


@dataclass(slots=True)
class ReviewerResult:
    grade: float | None
    grade_justification: str
    problems: list[ReviewerProblem]
    strengths: list[str]
    raw_response: str
    notes: list[str] = field(default_factory=list)

    @property
    def passed_threshold(self) -> float | None:
        return self.grade

    def to_report(self) -> str:
        lines = [f"# Independent reviewer report\n\n**Grade: {self.grade}**\n"]
        if self.grade_justification:
            lines.append(self.grade_justification + "\n")
        if self.problems:
            lines.append("## Top problems\n")
            for i, p in enumerate(self.problems, 1):
                lines.append(f"### {i}. [{p.priority}] {p.title}")
                lines.append(f'**Cytat:** "{p.quote}"\n')
                lines.append(f"**Dlaczego:** {p.why}\n")
                lines.append(f"**Naprawa:** {p.fix}\n")
        if self.strengths:
            lines.append("## Strengths\n")
            for s in self.strengths:
                lines.append(f"- {s}")
        if self.notes:
            lines.append("\n## Notes\n")
            lines.extend(f"- {n}" for n in self.notes)
        return "\n".join(lines)


def _extract_text(draft: Path, max_chars: int = 200_000) -> str:
    """Pull plain text from the docx. Truncate if oversized (Opus context limit guard)."""
    doc = Document(str(draft))
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[... pominięto pozostały tekst ze względu na limit kontekstu]"
    return text


def _parse_reviewer_response(text: str) -> ReviewerResult:
    grade: float | None = None
    m = re.search(r"Ocena[:\s]+(\d(?:[.,]\d)?)", text)
    if m:
        grade = float(m.group(1).replace(",", "."))

    justification = ""
    m = re.search(r"Ocena[:\s]+\d(?:[.,]\d)?\n(.+?)(?=\n##|\Z)", text, re.DOTALL)
    if m:
        justification = m.group(1).strip()

    problems = []
    for m in re.finditer(
        r"###\s*\d+\.\s*\[(PILNY|WYSOKI|ŚREDNI)\]\s*(.+?)\n"
        r"\*\*Cytat:\*\*\s*\"(.+?)\"\s*\n"
        r"\*\*Dlaczego[^:]*:\*\*\s*(.+?)\n"
        r"\*\*Naprawa:\*\*\s*(.+?)(?=\n###|\n##|\Z)",
        text,
        re.DOTALL,
    ):
        problems.append(ReviewerProblem(
            priority=m.group(1),
            title=m.group(2).strip(),
            quote=m.group(3).strip(),
            why=m.group(4).strip(),
            fix=m.group(5).strip(),
        ))

    strengths = []
    m = re.search(r"##\s*Mocne strony\s*\n(.+?)(?=\n##|\Z)", text, re.DOTALL)
    if m:
        strengths = [line.lstrip("- ").strip() for line in m.group(1).splitlines() if line.strip().startswith("-")]

    return ReviewerResult(
        grade=grade,
        grade_justification=justification,
        problems=problems,
        strengths=strengths,
        raw_response=text,
    )


def run_independent_review(project: ThesisProject, *, cli: ClaudeCLI | None = None) -> ReviewerResult:
    """Synchronous — one Opus call (via `claude -p`), no tool use needed."""
    draft = project.resolve_input(project.inputs.draft)
    thesis_text = _extract_text(draft)

    regulation_text = ""
    if project.inputs.regulation:
        try:
            reg = Document(str(project.resolve_input(project.inputs.regulation)))
            regulation_text = "\n".join(p.text for p in reg.paragraphs if p.text.strip())[:20_000]
        except Exception:
            pass

    user_msg = f"# Tekst pracy\n\n{thesis_text}"
    if regulation_text:
        user_msg += f"\n\n# Regulamin uczelni (formatowanie)\n\n{regulation_text}"

    cli = cli or ClaudeCLI()
    resp = cli.complete_sync(
        model=project.models.reviewer,
        system=REVIEWER_SYSTEM_PROMPT,
        user=user_msg,
        timeout=900,
    )
    result = _parse_reviewer_response(resp.text)
    if result.grade is None:
        result.notes.append("Could not parse grade from reviewer response — check raw_response.")
    return result


def save_report(result: ReviewerResult, project: ThesisProject) -> Path:
    out = project.reports_dir() / "independent_reviewer.md"
    out.write_text(result.to_report(), encoding="utf-8")
    # Also save raw response for audit trail
    raw = project.reports_dir() / "independent_reviewer_raw.txt"
    raw.write_text(result.raw_response, encoding="utf-8")
    return out
