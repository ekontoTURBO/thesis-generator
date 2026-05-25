"""The rolling progress report — disaster-recovery file.

The user said in the proven session:
    "te wszystkie plany informacje dodatki etapy pracy umieszczaj w jakimś
     temp markdownie na desktopie jakby się skończyło okno kontekstowe to
     żebyś wiedział co robić na czym jesteśmy"

So: at every phase boundary, this file is rewritten with the full state.
A fresh session reading it should be able to resume without any other context.

Format follows the proven session's template (brief section 6).
"""

from __future__ import annotations

import datetime as dt
import json
from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document

from thesis_generator.config import ThesisProject


@dataclass(slots=True)
class ProgressReport:
    phase: str
    status: str  # IN_PROGRESS | DONE | BLOCKED
    icon: str = "🎯"
    file_metrics: dict[str, str | int] = field(default_factory=dict)
    completed: list[str] = field(default_factory=list)
    findings: list[str] = field(default_factory=list)
    next_steps: list[str] = field(default_factory=list)
    blockers: list[str] = field(default_factory=list)
    timestamp: str = field(default_factory=lambda: dt.datetime.now().isoformat(timespec="seconds"))

    def render(self) -> str:
        lines = [f"# {self.icon} {self.phase} — {self.status}\n"]
        lines.append(f"_Updated: {self.timestamp}_\n")

        if self.file_metrics:
            lines.append("\n## Stan finalny\n")
            lines.append("| Element | Wartość |")
            lines.append("|---|---|")
            for k, v in self.file_metrics.items():
                lines.append(f"| {k} | {v} |")

        if self.completed:
            lines.append("\n## Co zostało zrobione\n")
            for i, item in enumerate(self.completed, 1):
                lines.append(f"{i}. {item}")

        if self.findings:
            lines.append("\n## Najważniejsze ustalenia\n")
            for f in self.findings:
                lines.append(f"- {f}")

        if self.next_steps:
            lines.append("\n## Następne kroki\n")
            for i, step in enumerate(self.next_steps, 1):
                lines.append(f"{i}. {step}")

        lines.append("\n## Blockers / decyzje do podjęcia\n")
        if self.blockers:
            for b in self.blockers:
                lines.append(f"- {b}")
        else:
            lines.append("- brak")
        return "\n".join(lines)


def measure_draft(draft: Path) -> dict[str, str | int]:
    """Pull the canonical metrics every progress report includes."""
    if not draft.exists():
        return {"draft": "(missing)"}
    doc = Document(str(draft))
    text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "Paragrafy": len(doc.paragraphs),
        "Tabele": len(doc.tables),
        "Obrazki": len(doc.inline_shapes),
        "Znaków": len(text),
        "Stron (~)": max(1, len(text) // 1800),
    }


def update_progress(
    project: ThesisProject,
    phase: str,
    status: str,
    *,
    icon: str = "🎯",
    completed: list[str] | None = None,
    findings: list[str] | None = None,
    next_steps: list[str] | None = None,
    blockers: list[str] | None = None,
    draft_override: Path | None = None,
) -> Path:
    """Append/replace the progress report on disk and return its path."""
    draft = draft_override or project.resolve_input(project.inputs.draft)
    metrics = measure_draft(draft)

    report = ProgressReport(
        phase=phase,
        status=status,
        icon=icon,
        file_metrics=metrics,
        completed=completed or [],
        findings=findings or [],
        next_steps=next_steps or [],
        blockers=blockers or [],
    )

    out = project.reports_dir() / "PROGRESS.md"
    out.write_text(report.render(), encoding="utf-8")

    # Also append to history.json for audit trail
    history = project.state_dir() / "progress_history.json"
    existing = []
    if history.exists():
        try:
            existing = json.loads(history.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass
    existing.append(asdict(report))
    history.write_text(json.dumps(existing, indent=2, ensure_ascii=False), encoding="utf-8")
    return out
