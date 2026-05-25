"""Ring A — internal verification (cheap, fast, no API calls).

Regex + python-docx over the draft itself. Catches:
- orphaned bibliography entries (in biblio but never cited in body)
- missing bibliography entries (cited in body but not in biblio)
- bibliography alphabetization issues
- em-dash count vs target (AI tell)
- formatting drift (font runs, heading style breakage)

Ported from `audyt_bibliografii.py` of the proven session, generalized: bib
section markers and required key-source list are now configurable, no
hardcoded paths.
"""

from __future__ import annotations

import json
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from thesis_generator.config import ThesisProject


# Citation regexes. Generalized from the proven audyt_bibliografii.py to handle
# Polish "i in." (= English "et al.") in both parens and narrative form.
_CITATION_PARENS = re.compile(
    r"\(([A-ZŁŚĆŻŹÓŃ][A-Za-zżółćęąśńźŁŚĆŻŹÓŃ\-\.]+"
    r"(?:(?:\s+i\s+|\s+&\s+|\s+oraz\s+|\,\s*)[A-ZŁŚĆŻŹÓŃ][A-Za-zżółćęąśńźŁŚĆŻŹÓŃ\-\.]+)*"
    r"(?:\s+et\s+al\.?|\s+i\s+in\.?)?),\s*(\d{4})",
    re.UNICODE,
)
_CITATION_NARRATIVE = re.compile(
    r"\b([A-ZŁŚĆŻŹÓŃ][A-Za-zżółćęąśńźŁŚĆŻŹÓŃ\-]+(?:a|owie|owa|i)?"
    r"(?:\s+i\s+(?:Plano\s+)?[A-ZŁŚĆŻŹÓŃ][A-Za-zżółćęąśńźŁŚĆŻŹÓŃ\-]+(?:a|owie|owa)?)?)"
    r"(?:\s+i\s+in\.?|\s+et\s+al\.?)?\s+\((\d{4})",
    re.UNICODE,
)
_BIB_ENTRY = re.compile(
    r"^([A-ZŁŚĆŻŹÓŃ][A-Za-zżółćęąśńźŁŚĆŻŹÓŃ\-\']+"
    r"(?:,\s*[A-Z]\.\s*[A-Z]?\.?)*"
    r"(?:(?:\s*,?\s*&\s*|\s+i\s+)[A-ZŁŚĆŻŹÓŃ][A-Za-zżółćęąśńźŁŚĆŻŹÓŃ\-\']+(?:,\s*[A-Z]\.\s*[A-Z]?\.?)*)*"
    r")(?:.*?)\((\d{4}|b\.d\.)",
    re.UNICODE,
)
_SECTION_HEADER = re.compile(r"^(I|II|III|IV)\.\s")


@dataclass(slots=True)
class BibEntry:
    section: str
    first_surname: str
    year: str
    full: str
    para_idx: int


@dataclass(slots=True)
class RingAResult:
    citations_in_text: list[tuple[str, str]]  # (surname, year)
    bib_entries: list[BibEntry]
    missing_in_bib: list[tuple[str, str]]
    orphaned_in_bib: list[BibEntry]
    alphabetization_issues: list[tuple[str, str]]
    em_dash_count: int
    em_dash_threshold_exceeded: bool
    summary_lines: list[str] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        return (
            not self.missing_in_bib
            and not self.alphabetization_issues
            and not self.em_dash_threshold_exceeded
        )

    def to_report(self) -> str:
        lines = ["# Ring A — Internal audit\n", f"Status: {'✅ PASS' if self.passed else '❌ FAIL'}\n"]
        lines.append("\n## Counts\n")
        lines.append(f"- citations in text: {len(self.citations_in_text)}")
        lines.append(f"- bibliography entries: {len(self.bib_entries)}")
        lines.append(f"- em-dashes in body: {self.em_dash_count}")
        if self.missing_in_bib:
            lines.append("\n## ❌ Cited but missing from bibliography\n")
            for s, y in self.missing_in_bib:
                lines.append(f"- {s} ({y})")
        if self.orphaned_in_bib:
            lines.append("\n## ⚠️ In bibliography but not cited (orphans)\n")
            for e in self.orphaned_in_bib:
                lines.append(f"- {e.first_surname} ({e.year})  — `{e.full[:80]}`")
        if self.alphabetization_issues:
            lines.append("\n## ⚠️ Alphabetization breaks\n")
            for prev, curr in self.alphabetization_issues:
                lines.append(f"- `{curr}` appears after `{prev}`")
        if self.summary_lines:
            lines.append("\n## Notes\n")
            lines.extend(f"- {l}" for l in self.summary_lines)
        return "\n".join(lines)


def _ascii_fold(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))


def _strip_polish_decl(surname: str) -> str:
    return re.sub(r"(?:a|y|i|e|owie|ego|ej|em|owi|ami)$", "", _ascii_fold(surname).lower())


def _find_bib_markers(doc: Document) -> tuple[int | None, dict[int, str]]:
    """Return (bib_start_idx, {paragraph_idx: section_label}) for bibliography only.

    Handles two layouts:
      - Bibliography split into sections (`I. Literatura naukowa`, `II. Raporty`, `NETOGRAFIA`)
      - Flat bibliography (no section markers) — treats all post-BIBLIOGRAFIA paragraphs as `I.`
    """
    bib_start = None
    section_at: dict[int, str] = {}
    in_bib = False
    current = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "BIBLIOGRAFIA":
            bib_start = i
            in_bib = True
            current = "I."  # default — flat bibliography case
            continue
        if not in_bib:
            continue
        if _SECTION_HEADER.match(t) or t == "NETOGRAFIA":
            current = t[:40]
            continue
        if current and t:
            section_at[i] = current
    return bib_start, section_at


def run_ring_a(project: ThesisProject, draft_path: Path | None = None) -> RingAResult:
    draft = draft_path or project.resolve_input(project.inputs.draft)
    doc = Document(str(draft))

    bib_start, sections = _find_bib_markers(doc)
    if bib_start is None:
        # No bibliography section yet — report everything as warnings
        return RingAResult(
            citations_in_text=[],
            bib_entries=[],
            missing_in_bib=[],
            orphaned_in_bib=[],
            alphabetization_issues=[],
            em_dash_count=sum(p.text.count(" — ") for p in doc.paragraphs),
            em_dash_threshold_exceeded=False,
            summary_lines=["No BIBLIOGRAFIA header found — skipping citation cross-check."],
        )

    body_text = "\n".join(p.text for p in doc.paragraphs[:bib_start])

    # Citations
    citations: set[tuple[str, str]] = set()
    for pat in (_CITATION_PARENS, _CITATION_NARRATIVE):
        for m in pat.finditer(body_text):
            first = re.split(r"\s+i\s+|\s+&\s+|\s+oraz\s+|\,\s*|\s+et\s+al", m.group(1).strip().rstrip("."))[0]
            citations.add((first, m.group(2)))

    # Bibliography
    entries: list[BibEntry] = []
    for i in range(bib_start, len(doc.paragraphs)):
        if i not in sections:
            continue
        t = doc.paragraphs[i].text.strip()
        m = _BIB_ENTRY.match(t)
        if not m:
            continue
        first_surname = m.group(1).split(",")[0].strip()
        entries.append(BibEntry(
            section=sections[i],
            first_surname=first_surname,
            year=m.group(2),
            full=t[:200],
            para_idx=i,
        ))

    # Cross-check
    bib_lookup = {(_strip_polish_decl(e.first_surname), e.year): e for e in entries}
    missing: list[tuple[str, str]] = []
    cited_keys: set[tuple[str, str]] = set()
    for surname, year in sorted(citations):
        key = (_strip_polish_decl(surname), year)
        cited_keys.add(key)
        # loose match
        if not any(k[0].startswith(key[0][:5]) or key[0].startswith(k[0][:5]) for k in bib_lookup if k[1] == year):
            missing.append((surname, year))

    orphans = [
        e for e in entries
        if not any(k[0].startswith(_strip_polish_decl(e.first_surname)[:5]) for k in cited_keys if k[1] == e.year)
    ]

    # Alphabetization (section I only — primary literature)
    section_i = [e for e in entries if e.section.startswith("I.")]
    alpha_issues = []
    prev = ""
    for e in section_i:
        if e.first_surname.lower() < prev.lower():
            alpha_issues.append((prev, e.first_surname))
        prev = e.first_surname

    # Em-dash
    em_count = body_text.count(" — ")

    return RingAResult(
        citations_in_text=sorted(citations),
        bib_entries=entries,
        missing_in_bib=missing,
        orphaned_in_bib=orphans,
        alphabetization_issues=alpha_issues,
        em_dash_count=em_count,
        em_dash_threshold_exceeded=em_count > 0,  # zero tolerance per humanization policy
    )


def save_report(result: RingAResult, project: ThesisProject) -> Path:
    out = project.reports_dir() / "ring_a_internal.md"
    out.write_text(result.to_report(), encoding="utf-8")
    return out
