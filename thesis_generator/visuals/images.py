"""Image insertion with caption + source — ports `insert_v75_images.py` pattern.

Each insertion produces three centered paragraphs:
  1. The image (centered, sized to the configured width in cm)
  2. The caption: "Rysunek N. Tytuł" (bold)
  3. The source: "Źródło: ..." (italic)
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


def _make_styled_caption_p(reference_para_element, text: str, *, bold: bool = False, italic: bool = False):
    """Build a centered paragraph with the requested run formatting,
    cloned from a reference paragraph so styles match the rest of the doc."""
    new_p = deepcopy(reference_para_element)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    ppr = new_p.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        new_p.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), "center")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    if len(rpr):
        run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    new_p.append(run)
    return new_p


def _pick_reference_paragraph(doc) -> object:
    """Find a long enough body paragraph to clone styles from."""
    for p in doc.paragraphs:
        if p.text.strip() and len(p.text) > 60:
            return p._element
    # Fallback: just use the first paragraph (even if empty)
    return doc.paragraphs[0]._element if doc.paragraphs else None


def insert_image_block(
    doc,
    *,
    after_paragraph,
    image_path: Path,
    caption: str,
    source: str,
    width_cm: float = 13.0,
) -> None:
    """Insert image + caption + source paragraphs immediately after `after_paragraph`.

    Caption is bold; source is italic. Both centered, cloned from a reference
    paragraph so they inherit the body font.
    """
    image_path = Path(image_path)
    if not image_path.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")

    ref_p_element = _pick_reference_paragraph(doc)

    # Image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    if after_paragraph is not None:
        after_paragraph._element.addnext(img_para._element)

    # Caption (bold)
    cap_p = _make_styled_caption_p(ref_p_element, caption, bold=True)
    img_para._element.addnext(cap_p)

    # Source (italic)
    src_p = _make_styled_caption_p(ref_p_element, source, italic=True)
    cap_p.addnext(src_p)
