"""Visuals pipeline — scan draft, build/insert each visual, append Spis tabel/rysunków.

Workflow:
  1. Parse all markers in the current draft.
  2. Assign final sequential numbers via `VisualsRegistry`.
  3. For each marker:
       - WYKRES: load data from xlsx → render PNG → insert as image block
       - TABELA with Dane: load xlsx range → create styled table after marker para
       - TABELA without Dane: assume writer already produced a raw table elsewhere; just register it
       - ILUSTRACJA: insert user-supplied image
       - SUGEROWANY: leave a placeholder note, skip insertion
  4. Strip markers from the body text.
  5. Style every existing table uniformly (catches manually-pasted tables too).
  6. Append Spis tabel + Spis rysunków at the end (before bibliography).
  7. Save and produce a markdown report.

Output: the draft is mutated in place (or to `output_path` if supplied). A report
lands at `_reports/visuals.md`.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from thesis_generator.config import ThesisProject
from thesis_generator.visuals.charts import (
    ChartData,
    load_chart_data_from_xlsx,
    render_chart,
)
from thesis_generator.visuals.images import insert_image_block
from thesis_generator.visuals.markers import (
    VisualKind,
    VisualMarker,
    parse_markers,
)
from thesis_generator.visuals.registry import VisualEntry, VisualsRegistry
from thesis_generator.visuals.tables import insert_table_from_xlsx, style_all_tables


@dataclass(slots=True)
class VisualsResult:
    inserted: int = 0
    skipped: list[tuple[str, str]] = field(default_factory=list)  # (marker title, reason)
    registry: VisualsRegistry | None = None
    output_path: Path | None = None
    notes: list[str] = field(default_factory=list)


def process_visuals(
    project: ThesisProject,
    *,
    output_path: Path | None = None,
) -> VisualsResult:
    """Run the full visuals pipeline against the project's draft."""
    draft = project.resolve_input(project.inputs.draft)
    target = Path(output_path) if output_path else draft

    doc = Document(str(draft))
    result = VisualsResult()
    registry = VisualsRegistry()
    result.registry = registry

    # Charts go here so they survive across reruns
    charts_dir = project.output_dir() / "_charts"
    charts_dir.mkdir(parents=True, exist_ok=True)

    # Walk every paragraph; for each, parse markers and process them in order.
    # We can't iterate `doc.paragraphs` while inserting tables/images after them
    # (it invalidates the iterator), so collect first then mutate.
    pending: list[tuple[int, VisualMarker]] = []
    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs):
        for m in parse_markers(p.text):
            pending.append((i, m))

    # Assign final numbers (in document order)
    entries = registry.assign([m for _, m in pending])

    # Process each marker
    for (para_idx, marker), entry in zip(pending, entries):
        anchor_para = paragraphs[para_idx]
        try:
            _process_one(project, doc, anchor_para, marker, entry, charts_dir, result)
        except Exception as e:
            result.skipped.append((marker.title, f"{type(e).__name__}: {e}"))

    # Strip markers from the body text
    for p in doc.paragraphs:
        for m in parse_markers(p.text):
            if m.raw_text in p.text:
                p.text = p.text.replace(m.raw_text, "").strip()

    # Style every table uniformly
    styled = style_all_tables(doc)
    result.notes.append(f"Styled {styled} tables uniformly.")

    # Append Spis tabel / Spis rysunków before BIBLIOGRAFIA (if present) or at end.
    _append_spisy(doc, registry, result)

    # Save
    doc.save(str(target))
    result.output_path = target

    # Persist report
    reports = project.reports_dir()
    (reports / "visuals.md").write_text(registry.render_full_report(), encoding="utf-8")

    return result


def _process_one(
    project: ThesisProject,
    doc,
    anchor_para,
    marker: VisualMarker,
    entry: VisualEntry,
    charts_dir: Path,
    result: VisualsResult,
) -> None:
    if marker.kind == VisualKind.SUGGESTED:
        result.skipped.append((marker.title, "suggested marker — needs [Dane: ...] spec"))
        return

    if marker.kind == VisualKind.IMAGE:
        if not marker.file:
            result.skipped.append((marker.title, "ILUSTRACJA without [Plik: ...]"))
            return
        image_path = project.resolve_input(Path(marker.file))
        if not image_path.exists():
            # Try inputs/visuals/ as a default
            fallback = project.project_dir / "inputs" / "visuals" / Path(marker.file).name
            if fallback.exists():
                image_path = fallback
            else:
                result.skipped.append((marker.title, f"image not found: {marker.file}"))
                return
        insert_image_block(
            doc,
            after_paragraph=anchor_para,
            image_path=image_path,
            caption=entry.caption,
            source=entry.source_line,
            width_cm=marker.width_cm or 13.0,
        )
        entry.image_path = str(image_path)
        result.inserted += 1
        return

    if marker.kind == VisualKind.CHART:
        if not marker.data:
            result.skipped.append((marker.title, "WYKRES without [Dane: ...]"))
            return
        data_file = project.resolve_input(Path(marker.data.file))
        if not data_file.exists():
            result.skipped.append((marker.title, f"data file not found: {marker.data.file}"))
            return
        cd = load_chart_data_from_xlsx(
            data_file, sheet=marker.data.sheet or "Sheet1", cells=marker.data.cells or "A1:B10"
        )
        cd.title = ""  # don't duplicate title in image; the caption carries it
        cd.chart_type = marker.chart_type or "bar"
        chart_filename = f"wykres_{entry.number:02d}_{_slugify(entry.title)}.png"
        chart_path = charts_dir / chart_filename
        render_chart(cd, chart_path, title=marker.title)
        insert_image_block(
            doc,
            after_paragraph=anchor_para,
            image_path=chart_path,
            caption=entry.caption,
            source=entry.source_line,
            width_cm=marker.width_cm or 14.0,
        )
        entry.image_path = str(chart_path)
        result.inserted += 1
        return

    if marker.kind == VisualKind.TABLE:
        if not marker.data:
            # No data spec → assume writer wrote a raw markdown/html table elsewhere;
            # we'll just register the entry, the user is on the hook for the body.
            result.skipped.append((marker.title, "TABELA without [Dane: ...] — only registered"))
            return
        data_file = project.resolve_input(Path(marker.data.file))
        if not data_file.exists():
            result.skipped.append((marker.title, f"data file not found: {marker.data.file}"))
            return
        # Insert caption BEFORE the table (Polish convention)
        from thesis_generator.visuals.images import _make_styled_caption_p, _pick_reference_paragraph
        ref_el = _pick_reference_paragraph(doc)
        cap_p = _make_styled_caption_p(ref_el, entry.caption, bold=True)
        anchor_para._element.addnext(cap_p)
        # Insert table after the caption
        insert_table_from_xlsx(
            doc,
            after_paragraph=anchor_para,
            file=data_file,
            sheet=marker.data.sheet or "Sheet1",
            cells=marker.data.cells or "A1:E10",
        )
        # Source line BELOW the table — we need to find the inserted table then put
        # a styled source paragraph after it. Simpler: append source paragraph after
        # the anchor (which currently has caption + table after it).
        src_p = _make_styled_caption_p(ref_el, entry.source_line, italic=True)
        # Walk siblings to find the table just inserted
        sibling = anchor_para._element.getnext()
        last_inserted = sibling
        for _ in range(3):  # caption, table, maybe more
            nxt = last_inserted.getnext() if last_inserted is not None else None
            if nxt is None:
                break
            last_inserted = nxt
        last_inserted.addnext(src_p)
        result.inserted += 1
        return


def _slugify(s: str, *, maxlen: int = 30) -> str:
    import re
    import unicodedata
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r"[^a-zA-Z0-9]+", "_", s).strip("_").lower()
    return s[:maxlen] or "x"


def _append_spisy(doc, registry: VisualsRegistry, result: VisualsResult) -> None:
    """Append Spis tabel + Spis rysunków before BIBLIOGRAFIA if present, else at end."""
    insertion_point = None
    for p in doc.paragraphs:
        if p.text.strip() == "BIBLIOGRAFIA":
            insertion_point = p
            break

    def _add_paragraph_block(title: str, entries: list[VisualEntry], label: str) -> None:
        if not entries:
            return
        heading = doc.add_paragraph(title)
        try:
            heading.style = doc.styles["Heading 1"]
        except KeyError:
            pass
        for e in entries:
            doc.add_paragraph(f"{label} {e.number}. {e.title}")

    # If we have an insertion point, build before it; else append at the end.
    # python-docx's add_paragraph always appends; we move elements afterwards.
    pre_count = len(doc.paragraphs)
    _add_paragraph_block("Spis tabel", registry.tables, "Tabela")
    _add_paragraph_block("Spis rysunków", registry.figures, "Rysunek")
    added = doc.paragraphs[pre_count:]

    if insertion_point is not None and added:
        # Move each added paragraph before the bibliography heading
        anchor = insertion_point._element
        for p in added:
            anchor.addprevious(p._element)

    if registry.tables:
        result.notes.append(f"Appended Spis tabel with {len(registry.tables)} entries.")
    if registry.figures:
        result.notes.append(f"Appended Spis rysunków with {len(registry.figures)} entries.")
