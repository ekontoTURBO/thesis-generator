"""Table styling + creation — ports `uniform_tables.py` from the proven session.

The proven style:
- Arial 10pt, black text
- White cells, D9D9D9 gray header
- Black single-line borders all around + insideH + insideV
- Header row bold
- Centered table alignment
- tblLook stripped so banded auto-formatting doesn't kick in

Two entry points:
- `style_all_tables(doc)` — apply the uniform style to every table already in
  a docx (use this after the writer drops raw tables in)
- `insert_table_from_xlsx(doc, after_paragraph, file, sheet, cells)` — read a
  rectangular range from Excel, create a styled table, insert it after the
  given paragraph
"""

from __future__ import annotations

from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HEADER_FILL = "D9D9D9"
BODY_FILL = "FFFFFF"
BLACK = RGBColor(0, 0, 0)


def _set_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    for old in list(tblPr.findall(qn("w:tblBorders"))):
        tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        b.append(el)
    tblPr.append(b)


def _shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _style_cell(cell, *, header: bool, font: str = "Arial", size_pt: int = 10) -> None:
    _shade(cell, HEADER_FILL if header else BODY_FILL)
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        if not p.runs and p.text == "":
            continue
        for r in p.runs:
            r.font.name = font
            r.font.color.rgb = BLACK
            if not r.font.size:
                r.font.size = Pt(size_pt)
            r.font.bold = bool(header)
            rpr = r._element.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.append(rf)
            for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rf.set(qn(a), font)


def style_table(tbl, *, font: str = "Arial", size_pt: int = 10) -> None:
    """Apply the uniform style to a single table in-place."""
    try:
        doc = tbl._parent
        if hasattr(doc, "styles"):
            tbl.style = doc.styles["Table Grid"]
    except Exception:
        pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.tblPr
    look = tblPr.find(qn("w:tblLook"))
    if look is not None:
        tblPr.remove(look)
    _set_borders(tbl)
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            _style_cell(cell, header=(ri == 0), font=font, size_pt=size_pt)


def style_all_tables(doc: Document, *, font: str = "Arial", size_pt: int = 10) -> int:
    """Apply the uniform style to every table in `doc`. Returns count styled."""
    n = 0
    for tbl in doc.tables:
        style_table(tbl, font=font, size_pt=size_pt)
        n += 1
    return n


def insert_table_from_xlsx(
    doc: Document,
    *,
    after_paragraph,
    file: Path,
    sheet: str,
    cells: str,
    has_header_row: bool = True,
) -> object:
    """Read a rectangular range from xlsx, create a styled table after the given paragraph.

    Returns the inserted `docx.table.Table`.
    """
    wb = openpyxl.load_workbook(str(file), data_only=True, read_only=True)
    ws = wb[sheet]
    rng = list(ws[cells])
    rows: list[list[str]] = []
    for row in rng:
        rows.append([_format_cell(c.value) for c in row])
    wb.close()

    if not rows:
        return None

    n_cols = len(rows[0])
    n_rows = len(rows)

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            tbl.cell(r_idx, c_idx).text = val
    style_table(tbl)

    # Move the new table to right after `after_paragraph`
    if after_paragraph is not None:
        after_paragraph._element.addnext(tbl._tbl)

    return tbl


def _format_cell(v) -> str:
    if v is None:
        return ""
    if isinstance(v, float):
        # If it's a whole number, drop the decimal; else 4-digit precision
        if v.is_integer():
            return str(int(v))
        return f"{v:.4f}".rstrip("0").rstrip(".")
    return str(v)
