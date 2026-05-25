"""Step 1 — Source inventory + style ingestion.

Builds two indices:
- `author_year → pdf_path` (so verifiers can find the source file for any citation)
- `xlsx_sheet → computed_stats` (so the writer doesn't re-run statistics already done)

Plus extracts the style fingerprint from the existing draft: heading style, citation
format detected, em-dash vs en-dash usage, chevron vs straight quotes, bibliography
structure. Future writes match the existing style — never invent new conventions.

Gotcha addressed (#11 — folder confusion): recursively maps EVERY PDF under EVERY
plausible root and emits a single index. No more "earlier agents looked in the
wrong folder and falsely declared sources missing."
"""

from __future__ import annotations

import json
import re
import unicodedata
from collections import Counter
from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document

from thesis_generator.config import ThesisProject


@dataclass(slots=True)
class SourceEntry:
    """One indexed source PDF."""

    path: str
    author_guess: str  # first author surname, stripped to ASCII for lookup
    year_guess: str | None
    size_bytes: int
    page_count_estimate: int | None = None


@dataclass(slots=True)
class StyleFingerprint:
    """Detected from the existing draft so new writes match it."""

    heading_styles_used: list[str] = field(default_factory=list)
    citation_pattern_examples: list[str] = field(default_factory=list)
    em_dash_count: int = 0
    en_dash_count: int = 0
    chevron_quote_count: int = 0  # «...»
    straight_quote_count: int = 0  # "..."
    paragraph_count: int = 0
    bibliography_sections: list[str] = field(default_factory=list)
    cited_pairs: list[tuple[str, str]] = field(default_factory=list)  # (surname, year)


@dataclass(slots=True)
class Inventory:
    sources: list[SourceEntry]
    style: StyleFingerprint
    notes: list[str]

    def to_dict(self) -> dict:
        return {
            "sources": [asdict(s) for s in self.sources],
            "style": asdict(self.style),
            "notes": self.notes,
        }


# ---------------------------------------------------------------- PDF indexing

_FILENAME_AUTHOR_YEAR = re.compile(
    r"^(?P<author>[A-ZŁŚĆŻŹÓŃ][A-Za-zŁłŚśĆćŻżŹźÓóŃń\-]+)"
    r"(?:[_,\s]+(?:et[\s_]*al\.?|[A-ZŁŚĆŻŹÓŃ][A-Za-zŁłŚśĆćŻżŹźÓóŃń\-]+))?"
    r".*?(?P<year>(?:19|20)\d{2}|b\.?d\.?)",
    re.UNICODE,
)


def _ascii_fold(s: str) -> str:
    """Strip Polish diacritics for lookup keys (Hocha → Hoch, Babbiego → Babbie)."""
    return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))


def _guess_author_year_from_filename(name: str) -> tuple[str, str | None]:
    stem = Path(name).stem
    m = _FILENAME_AUTHOR_YEAR.match(stem)
    if not m:
        # last resort: first all-caps-start word + any year-like substring
        words = re.findall(r"[A-ZŁŚĆŻŹÓŃ][A-Za-zŁłŚśĆćŻżŹźÓóŃń\-]+", stem)
        years = re.findall(r"(?:19|20)\d{2}", stem)
        return (words[0] if words else stem[:30], years[0] if years else None)
    author = _ascii_fold(m.group("author"))
    year = m.group("year")
    return author, year


def _index_pdfs(roots: list[Path]) -> list[SourceEntry]:
    seen: dict[str, SourceEntry] = {}
    for root in roots:
        if not root.exists():
            continue
        for p in root.rglob("*.pdf"):
            key = str(p.resolve())
            if key in seen:
                continue
            author, year = _guess_author_year_from_filename(p.name)
            seen[key] = SourceEntry(
                path=key,
                author_guess=author,
                year_guess=year,
                size_bytes=p.stat().st_size,
            )
    return sorted(seen.values(), key=lambda s: (s.author_guess.lower(), s.year_guess or ""))


# ------------------------------------------------------------- Style detection

_CITATION_INLINE = re.compile(
    r"\(([A-ZŁŚĆŻŹÓŃ][A-Za-zŁłŚśĆćŻżŹźÓóŃń\-\.]+"
    r"(?:(?:\s+i\s+|\s+&\s+|\,\s*)[A-ZŁŚĆŻŹÓŃ][A-Za-zŁłŚśĆćŻżŹźÓóŃń\-\.]+)*"
    r"(?:\s+et\s+al\.?)?),\s*(\d{4})",
    re.UNICODE,
)


def _extract_style(draft_path: Path) -> StyleFingerprint:
    doc = Document(str(draft_path))
    fp = StyleFingerprint()

    body_text = "\n".join(p.text for p in doc.paragraphs)
    fp.paragraph_count = len(doc.paragraphs)
    fp.em_dash_count = body_text.count(" — ")
    fp.en_dash_count = body_text.count(" – ")
    fp.chevron_quote_count = body_text.count("«") + body_text.count("»")
    fp.straight_quote_count = body_text.count('"')

    headings = Counter()
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            headings[p.style.name] += 1
    fp.heading_styles_used = [f"{name}: {n}" for name, n in headings.most_common()]

    # Citations + examples
    matches = list(_CITATION_INLINE.finditer(body_text))
    fp.citation_pattern_examples = [body_text[max(0, m.start() - 30) : m.end() + 30] for m in matches[:5]]
    fp.cited_pairs = sorted({(_ascii_fold(m.group(1).split()[0]), m.group(2)) for m in matches})

    # Bibliography section headers
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in ("BIBLIOGRAFIA", "NETOGRAFIA") or re.match(r"^(I|II|III|IV)\.\s", t):
            fp.bibliography_sections.append(t)

    return fp


# ------------------------------------------------------------- Main entrypoint


def build_inventory(project: ThesisProject) -> Inventory:
    """Top-level — runs both indexing passes and persists the result."""
    notes: list[str] = []

    # PDF index across all plausible roots (gotcha #11)
    roots = [project.resolve_input(project.inputs.sources_dir)]
    interviews = project.effective_interviews_dir()
    if interviews:
        roots.append(interviews)
    school = project.effective_school_dir()
    if school:
        roots.append(school)
    for extra in project.inputs.additional:
        roots.append(project.resolve_input(extra))
    sources = _index_pdfs(roots)
    notes.append(f"Indexed {len(sources)} unique PDFs across {len(roots)} root folder(s)")

    # Style from existing draft
    style = _extract_style(project.resolve_input(project.inputs.draft))
    notes.append(
        f"Draft has {style.paragraph_count} paragraphs, "
        f"{len(style.cited_pairs)} unique (author, year) citations, "
        f"{style.em_dash_count} em-dashes (AI tell if > 0 in final), "
        f"{style.chevron_quote_count} chevron-quote characters"
    )

    inv = Inventory(sources=sources, style=style, notes=notes)

    # Persist
    out = project.state_dir() / "inventory.json"
    out.write_text(json.dumps(inv.to_dict(), indent=2, ensure_ascii=False), encoding="utf-8")

    return inv


def load_inventory(project: ThesisProject) -> Inventory | None:
    out = project.state_dir() / "inventory.json"
    if not out.exists():
        return None
    data = json.loads(out.read_text(encoding="utf-8"))
    return Inventory(
        sources=[SourceEntry(**s) for s in data["sources"]],
        style=StyleFingerprint(**data["style"]),
        notes=data["notes"],
    )


def find_source_for_citation(inv: Inventory, surname: str, year: str | None) -> SourceEntry | None:
    """Lookup `author_year → pdf` with Polish-declension-aware matching."""
    target_author = _ascii_fold(surname).lower()
    root = re.sub(r"(?:a|y|i|e|owie|ego|ej|em|owi|ami)$", "", target_author)
    for s in inv.sources:
        cand = s.author_guess.lower()
        if cand.startswith(root[:5]) or target_author.startswith(cand[:5]):
            if year is None or s.year_guess == year:
                return s
    return None
